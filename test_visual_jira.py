"""Visual Jira Automation - See every step + Auto screenshot to Word doc"""
import asyncio
import sys
import os
from pathlib import Path

# Add src to path
sys.path.insert(0, 'src')

from tools.jira_browser_automation import JiraBrowserAutomation
from tools.document_parser import DocumentParser
from dotenv import load_dotenv
from datetime import datetime
import docx
from docx.shared import Inches

# Load environment variables
load_dotenv()

class VisualJiraWithScreenshots:
    """Visual Jira automation with automatic screenshots"""
    
    def __init__(self):
        self.jira_url = os.getenv('JIRA_URL')
        self.email = os.getenv('JIRA_EMAIL')
        self.api_token = os.getenv('JIRA_API_TOKEN')
        self.screenshots = []
        self.doc = None
        self.screenshot_dir = "screenshots"
        
        # Create screenshots directory
        Path(self.screenshot_dir).mkdir(exist_ok=True)
    
    def create_word_doc(self):
        """Create new Word document"""
        self.doc = docx.Document()
        self.doc.add_heading('Jira Automation - Visual Steps', 0)
        self.doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        self.doc.add_paragraph(f'Jira URL: {self.jira_url}')
        self.doc.add_page_break()
    
    def add_screenshot_to_doc(self, title: str, screenshot_path: str):
        """Add screenshot to Word doc with title"""
        if self.doc and Path(screenshot_path).exists():
            self.doc.add_heading(title, level=1)
            self.doc.add_paragraph(f'Timestamp: {datetime.now().strftime("%H:%M:%S")}')
            try:
                self.doc.add_picture(screenshot_path, width=Inches(6))
            except Exception as e:
                self.doc.add_paragraph(f'[Could not add image: {e}]')
            self.doc.add_page_break()
    
    async def run_with_screenshots(self, config: dict, take_screenshots: bool = True):
        """Run visual automation with screenshots"""
        
        print("\n" + "="*70)
        print("  🎬 VISUAL JIRA AUTOMATION WITH SCREENSHOTS")
        print("="*70)
        print(f"📸 Screenshots: {'ENABLED' if take_screenshots else 'DISABLED'}")
        print(f"📁 Screenshot folder: {self.screenshot_dir}/")
        print("="*70 + "\n")
        
        # Create Word doc if screenshots enabled
        if take_screenshots:
            self.create_word_doc()
        
        # Initialize browser automation
        automation = JiraBrowserAutomation(self.jira_url, self.email, self.api_token)
        
        # Start browser
        await automation.start_browser(headless=False)
        
        # Screenshot 1: Login
        print("🔐 Step 1: Logging into Jira...")
        await automation.login_to_jira()
        if take_screenshots:
            screenshot_path = f"{self.screenshot_dir}/01_login.png"
            await automation.page.screenshot(path=screenshot_path)
            self.add_screenshot_to_doc("1. Login to Jira", screenshot_path)
            print(f"  📸 Screenshot saved: {screenshot_path}")
        await automation.page.wait_for_timeout(2000)
        
        # Create Projects
        print("\n📁 Step 2: Creating Projects")
        print("-" * 70)
        
        for idx, proj in enumerate(config.get('projects', []), start=1):
            print(f"\n  Creating project {idx}: {proj['key']} - {proj['name']}")
            await automation.create_project_visual(proj['key'], proj['name'])
            
            if take_screenshots:
                screenshot_path = f"{self.screenshot_dir}/02_project_{proj['key']}.png"
                await automation.page.screenshot(path=screenshot_path)
                self.add_screenshot_to_doc(
                    f"2.{idx}. Project Created: {proj['key']}", 
                    screenshot_path
                )
                print(f"    📸 Screenshot saved: {screenshot_path}")
            
            await automation.page.wait_for_timeout(2000)
        
        # Create Epics
        print("\n📋 Step 3: Creating Epics")
        print("-" * 70)
        
        for idx, epic in enumerate(config.get('epics', []), start=1):
            print(f"\n  Creating epic {idx}: {epic['name']}")
            await automation.create_epic_visual(epic['project'], epic['name'])
            
            if take_screenshots:
                screenshot_path = f"{self.screenshot_dir}/03_epic_{idx}.png"
                await automation.page.screenshot(path=screenshot_path)
                self.add_screenshot_to_doc(
                    f"3.{idx}. Epic Created: {epic['name']}", 
                    screenshot_path
                )
                print(f"    📸 Screenshot saved: {screenshot_path}")
            
            await automation.page.wait_for_timeout(1500)
        
        # Create Stories
        print("\n📝 Step 4: Creating Stories")
        print("-" * 70)
        
        for idx, story in enumerate(config.get('stories', []), start=1):
            print(f"\n  Creating story {idx}: {story['summary']}")
            await automation.create_story_visual(
                story['project'],
                story['summary'],
                epic_name=story.get('epic_id'),
                labels=story.get('labels'),
                story_points=story.get('story_points')
            )
            
            if take_screenshots:
                screenshot_path = f"{self.screenshot_dir}/04_story_{idx}.png"
                await automation.page.screenshot(path=screenshot_path)
                self.add_screenshot_to_doc(
                    f"4.{idx}. Story Created: {story['summary'][:50]}", 
                    screenshot_path
                )
                print(f"    📸 Screenshot saved: {screenshot_path}")
            
            await automation.page.wait_for_timeout(1000)
        
        # Final screenshot
        if take_screenshots:
            screenshot_path = f"{self.screenshot_dir}/05_final.png"
            await automation.page.screenshot(path=screenshot_path)
            self.add_screenshot_to_doc("5. Final View - All Items Created", screenshot_path)
            print(f"\n  📸 Final screenshot saved: {screenshot_path}")
        
        # Save Word document
        if take_screenshots:
            doc_path = f"jira_automation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            self.doc.save(doc_path)
            print(f"\n📄 Word document saved: {doc_path}")
        
        print("\n" + "="*70)
        print("  ✅ AUTOMATION COMPLETE!")
        print("="*70)
        print("\n👀 Browser will stay open for 30 seconds...")
        await automation.page.wait_for_timeout(30000)
        
        # Close browser
        if automation.browser:
            await automation.browser.close()
        if automation.playwright:
            await automation.playwright.stop()
        
        return {"success": True, "screenshots": len(self.screenshots) if take_screenshots else 0}


async def main():
    """Main function"""
    
    print("\n🎬 Visual Jira Automation Tool")
    print("="*70)
    print("Options:")
    print("  1. Run with screenshots (creates Word doc)")
    print("  2. Run without screenshots (just watch)")
    print("="*70)
    
    choice = input("\nChoice (1 or 2): ").strip()
    take_screenshots = choice == "1"
    
    # Parse document
    print("\n📄 Step 1: Parsing assignment document...")
    parser = DocumentParser()
    
    doc_path = 'documents/Assignment-11_Monday.docx'
    if not Path(doc_path).exists():
        print(f"❌ Document not found: {doc_path}")
        print("Please check the path and try again.")
        return
    
    result = await parser.parse_document(doc_path, use_ai=False)
    
    if not result['success']:
        print(f"❌ Parse failed: {result['error']}")
        return
    
    config = result['data']
    print(f"✅ Parsed successfully!")
    print(f"   Projects: {len(config.get('projects', []))}")
    print(f"   Epics: {len(config.get('epics', []))}")
    print(f"   Stories: {len(config.get('stories', []))}")
    
    # Run visual automation
    automation = VisualJiraWithScreenshots()
    result = await automation.run_with_screenshots(config, take_screenshots)
    
    if result['success']:
        print("\n✅ ALL DONE!")
        if take_screenshots:
            print(f"📸 {result['screenshots']} screenshots captured")
            print(f"📄 Word document created with all screenshots")
    else:
        print(f"\n❌ Failed: {result.get('error')}")


if __name__ == "__main__":
    asyncio.run(main())