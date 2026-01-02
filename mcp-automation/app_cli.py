# #!/usr/bin/env python3
# """CLI Application - Main Entry Point"""
# import asyncio
# import os
# import sys
# from src.config import Config
# from src.mcp_server import MCPServer
# from src.ai_client import AIClient

# async def main():
#     print("=" * 60)
#     print("  ⚡ MCP AI Automation System - Groq AI")
#     print("  🚀 Ultra-Fast AI Inference")
#     print("=" * 60)
#     print()
    
#     # Validate config
#     try:
#         Config.validate()
#         print("✅ Groq API key loaded")
#     except ValueError as e:
#         print(f"❌ Configuration error: {e}")
#         print("\nPlease create a .env file with:")
#         print("GROQ_API_KEY=gsk_your-key-here")
#         print("\nGet your key from: https://console.groq.com/")
#         return
    
#     # Initialize
#     print("🔧 Initializing MCP server...")
#     mcp = MCPServer()
    
#     if Config.AWS_ACCESS_KEY_ID and Config.AWS_SECRET_ACCESS_KEY:
#         mcp.configure_aws(
#             Config.AWS_ACCESS_KEY_ID,
#             Config.AWS_SECRET_ACCESS_KEY,
#             Config.AWS_REGION
#         )
#         print(f"✅ AWS configured (region: {Config.AWS_REGION})")
#     else:
#         print("⚠️  AWS not configured (optional)")
    
#     print("⚡ Initializing AI client with Groq (ultra-fast)...")
#     ai = AIClient(Config.GROQ_API_KEY, mcp)
#     print("✅ System ready!")
#     print()
    
#     print("💡 Example commands:")
#     print("   • play music on youtube")
#     print("   • create s3 bucket test-bucket-2024")
#     print("   • go to google.com")
#     print("   • list files")
#     print()
#     print("💨 Groq is 10x faster than OpenAI!")
#     print("Commands: 'quit' or 'exit' to stop, 'history' for past tasks")
#     print("=" * 60)
#     print()
    
#     while True:
#         try:
#             command = input("⚡ Groq Command: ").strip()
            
#             if not command:
#                 continue
            
#             if command.lower() in ['quit', 'exit', 'q']:
#                 print("\n👋 Goodbye!")
#                 await mcp.browser.close()
#                 break
            
#             if command.lower() == 'history':
#                 print("\n📋 Recent Tasks:")
#                 print("-" * 60)
#                 tasks = mcp.db.get_recent_tasks(10)
#                 for task in tasks:
#                     status_icon = "✅" if task[3] == "completed" else "❌"
#                     print(f"{status_icon} [{task[6][:19]}] {task[1][:50]}")
#                 print("-" * 60)
#                 print()
#                 continue
            
#             print(f"\n▶️  Executing: {command}")
#             print("⚡ Groq is processing at lightning speed...\n")
            
#             result = await ai.execute(command)
            
#             print("\n" + "=" * 60)
#             if result["success"]:
#                 print("✅ Success!")
#                 print(f"\n{result['result']['message']}")
#                 print(f"\n📊 AI Iterations: {result['result']['iterations']}")
#                 print(f"⏱️  Time: {result['execution_time']:.2f}s ⚡")
#             else:
#                 print("❌ Failed!")
#                 print(f"\nError: {result['error']}")
#             print("=" * 60)
#             print()
            
#         except KeyboardInterrupt:
#             print("\n\n👋 Interrupted!")
#             await mcp.browser.close()
#             break
#         except Exception as e:
#             print(f"\n❌ Error: {e}\n")

# if __name__ == "__main__":
#     try:
#         asyncio.run(main())
#     except KeyboardInterrupt:
#         print("\n\n👋 Goodbye!")

#!/usr/bin/env python3

"""CLI Application - Main Entry Point - COMPLETE VERSION with all features"""
import asyncio
import os
import sys
from pathlib import Path
from datetime import datetime
from src.config import Config
from src.mcp_server import MCPServer
from src.ai_client import AIClient
import docx
from docx.shared import Inches

async def handle_jira_visual_command(mcp, command):
    """Handle Jira visual commands with screenshots"""
    
    # Extract document path from command
    doc_path = None
    words = command.split()
    for word in words:
        if 'documents/' in word or 'documents\\' in word or '.docx' in word or '.pdf' in word:
            doc_path = word.strip('"\'')
            break
    
    # Default path if not specified
    if not doc_path:
        doc_path = 'documents/Assignment-11 Monday.docx'
    
    # Convert to Path object (handles both / and \)
    doc_path = Path(doc_path)
    
    # Check if file exists
    if not doc_path.exists():
        print(f"❌ File not found: {doc_path}")
        print("💡 Available files:")
        docs_dir = Path('documents')
        if docs_dir.exists():
            for f in docs_dir.glob('*.docx'):
                print(f"   - {f}")
            print("\n💡 Try one of these commands:")
            for f in docs_dir.glob('*.docx'):
                print(f'   complete jira assignment "{f}"')
        return
    
    # Check if screenshots requested
    with_screenshots = any(word in command.lower() for word in ['screenshot', 'screenshots', 'ss', 'visual', 'document'])
    
    print(f"\n🎬 Starting Visual Jira Automation")
    if with_screenshots:
        print("📸 WITH Screenshots - Word document will be generated")
    print(f"📄 Document: {doc_path}")
    print("="*60)
    
    # Parse document
    from src.tools.document_parser import DocumentParser
    parser = DocumentParser()
    parse_result = await parser.parse_document(str(doc_path), use_ai=False)
    
    if not parse_result['success']:
        print(f"❌ Parse failed: {parse_result['error']}")
        return
    
    config = parse_result['data']
    print(f"✅ Parsed: {len(config.get('projects', []))} projects, {len(config.get('epics', []))} epics, {len(config.get('stories', []))} stories\n")
    
    if with_screenshots:
        # Create with screenshots
        screenshot_dir = Path("screenshots")
        screenshot_dir.mkdir(exist_ok=True)
        
        # Initialize Word doc
        screenshot_doc = docx.Document()
        screenshot_doc.add_heading('Jira Automation - Visual Steps', 0)
        screenshot_doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        screenshot_doc.add_paragraph(f'Document: {doc_path}')
        screenshot_doc.add_page_break()
        
        # Start browser
        await mcp.jira_visual.start_browser(headless=False)
        
        # Login (manual)
        await mcp.jira_visual.login_to_jira()
        if not mcp.jira_visual.is_logged_in:
            print("❌ Login failed")
            return
        
        # Screenshot after login
        screenshot_path = screenshot_dir / "01_login.png"
        await mcp.jira_visual.page.screenshot(path=str(screenshot_path))
        screenshot_doc.add_heading("1. Login to Jira", level=1)
        screenshot_doc.add_picture(str(screenshot_path), width=Inches(6))
        screenshot_doc.add_page_break()
        print("  📸 Screenshot: Login")
        
        # Create Projects
        print("\n📁 Creating Projects...")
        for idx, proj in enumerate(config.get('projects', []), start=1):
            await mcp.jira_visual.create_project_visual(proj['key'], proj['name'])
            screenshot_path = screenshot_dir / f"02_project_{proj['key']}.png"
            await mcp.jira_visual.page.screenshot(path=str(screenshot_path))
            screenshot_doc.add_heading(f"2.{idx}. Project: {proj['key']}", level=1)
            screenshot_doc.add_picture(str(screenshot_path), width=Inches(6))
            screenshot_doc.add_page_break()
            print(f"  📸 Screenshot: Project {proj['key']}")
            await mcp.jira_visual.page.wait_for_timeout(2000)
        
        # Create Epics
        print("\n📋 Creating Epics...")
        for idx, epic in enumerate(config.get('epics', []), start=1):
            await mcp.jira_visual.create_epic_visual(epic['project'], epic['name'])
            screenshot_path = screenshot_dir / f"03_epic_{idx}.png"
            await mcp.jira_visual.page.screenshot(path=str(screenshot_path))
            screenshot_doc.add_heading(f"3.{idx}. Epic: {epic['name'][:50]}", level=1)
            screenshot_doc.add_picture(str(screenshot_path), width=Inches(6))
            screenshot_doc.add_page_break()
            print(f"  📸 Screenshot: Epic {idx}")
            await mcp.jira_visual.page.wait_for_timeout(1500)
        
        # Create Stories
        print("\n📝 Creating Stories...")
        for idx, story in enumerate(config.get('stories', []), start=1):
            await mcp.jira_visual.create_story_visual(story['project'], story['summary'])
            screenshot_path = screenshot_dir / f"04_story_{idx}.png"
            await mcp.jira_visual.page.screenshot(path=str(screenshot_path))
            screenshot_doc.add_heading(f"4.{idx}. Story: {story['summary'][:50]}", level=1)
            screenshot_doc.add_picture(str(screenshot_path), width=Inches(6))
            screenshot_doc.add_page_break()
            print(f"  📸 Screenshot: Story {idx}")
            await mcp.jira_visual.page.wait_for_timeout(1000)
        
        # Final screenshot
        screenshot_path = screenshot_dir / "05_final.png"
        await mcp.jira_visual.page.screenshot(path=str(screenshot_path))
        screenshot_doc.add_heading("5. Final View", level=1)
        screenshot_doc.add_picture(str(screenshot_path), width=Inches(6))
        print("  📸 Screenshot: Final view")
        
        # Save Word doc
        doc_filename = f"jira_automation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        screenshot_doc.save(doc_filename)
        
        print(f"\n✅ Word document saved: {doc_filename}")
        print("👀 Browser will stay open for 30 seconds...")
        await mcp.jira_visual.page.wait_for_timeout(30000)
        
        # Close browser
        try:
            if mcp.jira_visual.browser and mcp.jira_visual.browser.is_connected():
                await mcp.jira_visual.browser.close()
            if mcp.jira_visual.playwright:
                await mcp.jira_visual.playwright.stop()
        except:
            pass
    else:
        # Without screenshots (just visual)
        result = await mcp.jira_visual.create_assignment_visual(config)
        if result['success']:
            print(f"\n✅ {result['message']}")
        else:
            print(f"\n❌ Failed: {result.get('error')}")


async def main():
    print("=" * 60)
    print("  ⚡ MCP AI Automation System - Groq AI")
    print("  🚀 Ultra-Fast AI Inference + Advanced Features")
    print("=" * 60)
    print()
    
    # Validate config
    try:
        Config.validate()
        print("✅ Groq API key loaded")
    except ValueError as e:
        print(f"❌ Configuration error: {e}")
        print("\nPlease create a .env file with:")
        print("GROQ_API_KEY=gsk_your-key-here")
        print("\nGet your key from: https://console.groq.com/")
        return
    
    # Initialize
    print("🔧 Initializing MCP server...")
    mcp = MCPServer()
    
    # AWS configuration
    if Config.is_aws_configured():
        mcp.configure_aws(
            Config.AWS_ACCESS_KEY_ID,
            Config.AWS_SECRET_ACCESS_KEY,
            Config.AWS_REGION
        )
        print(f"✅ AWS configured (region: {Config.AWS_REGION})")
    else:
        print("⚠️  AWS not configured (optional)")
    
    # Jira configuration
    if Config.is_jira_configured():
        mcp.configure_jira(
            Config.JIRA_URL,
            Config.JIRA_EMAIL,
            Config.JIRA_API_TOKEN
        )
        print(f"✅ Jira configured ({Config.JIRA_URL})")
    else:
        print("⚠️  Jira not configured (optional)")
    
    # GitHub configuration - WITH DEBUG
    print("\n🔍 GitHub Configuration Check:")
    print(f"   Config.GITHUB_TOKEN exists: {bool(Config.GITHUB_TOKEN)}")
    
    if Config.is_github_configured():
        print(f"   ✅ Token found, configuring GitHub...")
        mcp.configure_github(Config.GITHUB_TOKEN)
        
        # CHECK if it actually worked
        if mcp.github is not None:
            print(f"   ✅ GitHub configured successfully")
            print(f"   ✅ mcp.github = {type(mcp.github)}")
        else:
            print(f"   ❌ GitHub configuration FAILED!")
            print(f"   ❌ mcp.github is still None")
            print(f"   ❌ This means PyGithub is NOT installed!")
            print(f"\n   🔧 FIX: Run this command:")
            print(f"      pip install PyGithub")
    else:
        print("   ⚠️  GitHub not configured (GITHUB_TOKEN missing in .env)")
    
    # Vercel configuration
    if Config.is_vercel_configured():
        mcp.configure_vercel(Config.VERCEL_TOKEN)
        print(f"✅ Vercel configured")
    else:
        print("⚠️  Vercel not configured (optional - add VERCEL_TOKEN to .env)")
    
    # REMOVED: Figma configuration - now handled automatically in mcp_server.py
    # The figma_to_website tool will use FIGMA_TOKEN from .env directly
    figma_token = Config.FIGMA_TOKEN if hasattr(Config, 'FIGMA_TOKEN') else None
    if figma_token:
        print(f"✅ Figma token configured (will be used by figma_to_website tool)")
    else:
        print("⚠️  Figma token not set (optional - add FIGMA_TOKEN to .env)")
    
    # Initialize AI client
    print("⚡ Initializing AI client with Groq (ultra-fast)...")
    ai = AIClient(Config.GROQ_API_KEY, mcp)
    
    # Document Generator configuration
    mcp.configure_document_generator(ai)
    print(f"✅ Document Generator configured")
    
    print("✅ System ready!")
    print()
    
    print("💡 Example commands:")
    print("   🌐 Browser:")
    print("      • play music on youtube")
    print("      • go to google.com")
    print()
    print("   🐙 GitHub:")
    print("      • create github repo my-awesome-project")
    print("      • build this project and push to github as mcp-automation")
    print("      • generate readme for this project")
    print()
    print("   ▲ Vercel:")
    print("      • deploy this on vercel")
    print("      • list my vercel deployments")
    print()
    print("   🎨 Figma to Website (COMPLETE WORKFLOW):")
    print("      • build website from figma [URL] project: my-site")
    print("      • convert figma design to website [URL]")
    print("      👉 Auto-deploys to GitHub + Vercel!")
    print()
    print("   📄 Document Generator:")
    print("      • create project proposal for fitness app")
    print("      • generate business plan for AI startup")
    print("      • make presentation about machine learning")
    print()
    print("   📋 Jira:")
    print('      • complete jira assignment "documents\\Assignment.docx"')
    print("      • complete jira assignment with screenshots")
    print()
    print("   📁 Files:")
    print("      • list files")
    print("      • read file config.py")
    print()
    print("💨 Groq is 10x faster than OpenAI!")
    print("Commands: 'quit' or 'exit' to stop, 'history' for past tasks")
    print("=" * 60)
    print()
    
    while True:
        try:
            command = input("⚡ Command: ").strip()
            
            if not command:
                continue
            
            if command.lower() in ['quit', 'exit', 'q']:
                print("\n👋 Goodbye!")
                await mcp.browser.close()
                break
            
            if command.lower() == 'history':
                print("\n📋 Recent Tasks:")
                print("-" * 60)
                tasks = mcp.db.get_recent_tasks(10)
                for task in tasks:
                    status_icon = "✅" if task[3] == "completed" else "❌"
                    print(f"{status_icon} [{task[6][:19]}] {task[1][:50]}")
                print("-" * 60)
                print()
                continue
            
            # Check for Jira commands
            if any(word in command.lower() for word in ['jira', 'assignment', 'complete']):
                await handle_jira_visual_command(mcp, command)
                continue
            
            print(f"\n▶️  Executing: {command}")
            print("⚡ Groq is processing at lightning speed...\n")
            
            result = await ai.execute(command)
            
            print("\n" + "=" * 60)
            if result["success"]:
                print("✅ Success!")
                print(f"\n{result['result']['message']}")
                print(f"\n📊 AI Iterations: {result['result']['iterations']}")
                print(f"⏱️  Time: {result['execution_time']:.2f}s ⚡")
            else:
                print("❌ Failed!")
                print(f"\nError: {result['error']}")
            print("=" * 60)
            print()
            
        except KeyboardInterrupt:
            print("\n\n👋 Interrupted!")
            await mcp.browser.close()
            break
        except Exception as e:
            print(f"\n❌ Error: {e}\n")

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n\n👋 Goodbye!")