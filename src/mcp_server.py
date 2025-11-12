# """MCP Server implementation - Fixed for MCP 1.19.0+"""
# from mcp.server import Server
# from mcp.types import Tool, TextContent
# from .tools import BrowserTool, AWSTool, SystemTool
# from .database import DatabaseManager
# import logging

# logger = logging.getLogger(__name__)

# class MCPServer:
#     """MCP Server managing automation tools"""
    
#     def __init__(self):
#         self.server = Server("automation-server")
#         self.browser = BrowserTool()
#         self.aws = None
#         self.system = SystemTool()
#         self.db = DatabaseManager()
        
#         # Store tools list manually
#         self.tools_list = []
        
#         self._register_tools()
    
#     def configure_aws(self, access_key: str, secret_key: str, region: str):
#         """Configure AWS"""
#         self.aws = AWSTool(access_key, secret_key, region)
#         logger.info(f"AWS configured: {region}")
    
#     def _register_tools(self):
#         """Register MCP tools"""
        
#         # Define all tools
#         self.tools_list = [
#             Tool(
#                 name="browser_navigate",
#                 description="Navigate to a URL in the browser",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "url": {"type": "string", "description": "URL to navigate to"}
#                     },
#                     "required": ["url"]
#                 }
#             ),
#             Tool(
#                 name="youtube_play",
#                 description="Search and play a video on YouTube",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "query": {"type": "string", "description": "Video search query"}
#                     },
#                     "required": ["query"]
#                 }
#             ),
#             Tool(
#                 name="browser_screenshot",
#                 description="Take a screenshot of the current page",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "filename": {
#                             "type": "string", 
#                             "description": "Optional filename for screenshot"
#                         }
#                     }
#                 }
#             ),
#             Tool(
#                 name="aws_create_bucket",
#                 description="Create an AWS S3 bucket",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "bucket_name": {
#                             "type": "string",
#                             "description": "Name of the bucket to create"
#                         }
#                     },
#                     "required": ["bucket_name"]
#                 }
#             ),
#             Tool(
#                 name="aws_list_buckets",
#                 description="List all AWS S3 buckets",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {}
#                 }
#             ),
#             Tool(
#                 name="file_read",
#                 description="Read contents of a file",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "path": {"type": "string", "description": "Path to the file"}
#                     },
#                     "required": ["path"]
#                 }
#             ),
#             Tool(
#                 name="file_write",
#                 description="Write content to a file",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "path": {"type": "string", "description": "Path to the file"},
#                         "content": {"type": "string", "description": "Content to write"}
#                     },
#                     "required": ["path", "content"]
#                 }
#             ),
#             Tool(
#                 name="list_directory",
#                 description="List contents of a directory",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "path": {
#                             "type": "string",
#                             "description": "Directory path (default: current directory)"
#                         }
#                     }
#                 }
#             ),
#         ]
        
#         logger.info(f"Registered {len(self.tools_list)} MCP tools")
        
#         # Register handlers with MCP server
#         @self.server.list_tools()
#         async def list_tools() -> list[Tool]:
#             """Return list of available tools"""
#             return self.tools_list
        
#         @self.server.call_tool()
#         async def call_tool(name: str, arguments: dict) -> list[TextContent]:
#             """Execute a tool"""
#             logger.info(f"Tool called: {name} with args: {arguments}")
            
#             try:
#                 result = None
                
#                 # Browser tools
#                 if name == "browser_navigate":
#                     result = await self.browser.navigate(arguments["url"])
                    
#                 elif name == "youtube_play":
#                     result = await self.browser.play_youtube(arguments["query"])
                    
#                 elif name == "browser_screenshot":
#                     filename = arguments.get("filename", "screenshot.png")
#                     result = await self.browser.screenshot(filename)
                
#                 # AWS tools
#                 elif name == "aws_create_bucket":
#                     if not self.aws:
#                         result = {"success": False, "error": "AWS not configured"}
#                     else:
#                         result = await self.aws.create_bucket(arguments["bucket_name"])
                        
#                 elif name == "aws_list_buckets":
#                     if not self.aws:
#                         result = {"success": False, "error": "AWS not configured"}
#                     else:
#                         result = await self.aws.list_buckets()
                
#                 # System tools
#                 elif name == "file_read":
#                     result = await self.system.read_file(arguments["path"])
                    
#                 elif name == "file_write":
#                     result = await self.system.write_file(
#                         arguments["path"],
#                         arguments["content"]
#                     )
                    
#                 elif name == "list_directory":
#                     path = arguments.get("path", ".")
#                     result = await self.system.list_directory(path)
                
#                 else:
#                     result = {"success": False, "error": f"Unknown tool: {name}"}
                
#                 # Log tool usage
#                 self.db.log_tool_usage(name, result.get("success", False), 0)
                
#                 # Format response
#                 if result.get("success"):
#                     message = f"✅ {result.get('message', 'Success')}"
#                     if result.get('data'):
#                         # Add relevant data to message
#                         if 'title' in result['data']:
#                             message += f"\nTitle: {result['data']['title']}"
#                         if 'buckets' in result['data']:
#                             message += f"\nBuckets: {', '.join(result['data']['buckets'][:5])}"
#                         if 'files' in result['data']:
#                             message += f"\nFiles: {len(result['data']['files'])}"
#                         if 'content' in result['data']:
#                             content = result['data']['content'][:200]
#                             message += f"\nContent: {content}..."
#                 else:
#                     message = f"❌ Error: {result.get('error', 'Unknown error')}"
                
#                 return [TextContent(type="text", text=message)]
                
#             except Exception as e:
#                 logger.error(f"Tool execution error: {e}", exc_info=True)
#                 return [TextContent(
#                     type="text",
#                     text=f"❌ Tool execution failed: {str(e)}"
#                 )]
    
#     async def get_tools_list(self):
#         """Get list of available tools"""
#         return self.tools_list
    
#     async def execute_tool(self, name: str, arguments: dict):
#         """Execute a tool directly"""
#         logger.info(f"Executing tool: {name} with args: {arguments}")
        
#         try:
#             result = None
            
#             # Browser tools
#             if name == "browser_navigate":
#                 result = await self.browser.navigate(arguments["url"])
                
#             elif name == "youtube_play":
#                 result = await self.browser.play_youtube(arguments["query"])
                
#             elif name == "browser_screenshot":
#                 filename = arguments.get("filename", "screenshot.png")
#                 result = await self.browser.screenshot(filename)
            
#             # AWS tools
#             elif name == "aws_create_bucket":
#                 if not self.aws:
#                     result = {"success": False, "error": "AWS not configured"}
#                 else:
#                     result = await self.aws.create_bucket(arguments["bucket_name"])
                    
#             elif name == "aws_list_buckets":
#                 if not self.aws:
#                     result = {"success": False, "error": "AWS not configured"}
#                 else:
#                     result = await self.aws.list_buckets()
            
#             # System tools
#             elif name == "file_read":
#                 result = await self.system.read_file(arguments["path"])
                
#             elif name == "file_write":
#                 result = await self.system.write_file(
#                     arguments["path"],
#                     arguments["content"]
#                 )
                
#             elif name == "list_directory":
#                 path = arguments.get("path", ".")
#                 result = await self.system.list_directory(path)
            
#             else:
#                 result = {"success": False, "error": f"Unknown tool: {name}"}
            
#             # Log usage
#             self.db.log_tool_usage(name, result.get("success", False), 0)
            
#             return result
            
#         except Exception as e:
#             logger.error(f"Tool execution error: {e}", exc_info=True)
#             return {"success": False, "error": str(e)}

"""MCP Server implementation - Updated with GitHub and Vercel"""
from mcp.server import Server
from mcp.types import Tool, TextContent
from .tools import BrowserTool, AWSTool, SystemTool
try:
    from .tools import UniversalJiraAutomation
except ImportError:
    UniversalJiraAutomation = None

try:
    from .tools import GitHubTool
except ImportError:
    GitHubTool = None

try:
    from .tools import VercelTool
except ImportError:
    VercelTool = None
try:
    from .tools import FigmaToWebsiteTool
except ImportError:
    FigmaToWebsiteTool = None

try:
    from .tools import DocumentGeneratorTool
except ImportError:
    DocumentGeneratorTool = None

from .database import DatabaseManager
import logging
from pathlib import Path
from datetime import datetime
import docx
from docx.shared import Inches

logger = logging.getLogger(__name__)

class MCPServer:
    """MCP Server managing automation tools"""
    
    def __init__(self):
        self.server = Server("automation-server")
        self.browser = BrowserTool()
        self.aws = None
        self.system = SystemTool()
        self.jira = None
        self.jira_visual = None  # Visual automation
        self.github = None  # GitHub tool
        self.vercel = None  # Vercel tool
        self.figma = None  # Figma to Website tool
        self.doc_gen = None  # Document generator tool
        self.db = DatabaseManager()
        
        # Screenshot management
        self.screenshot_enabled = False
        self.screenshot_doc = None
        self.screenshot_dir = "screenshots"
        
        # Store tools list manually
        self.tools_list = []
        
        self._register_tools()
    
    def configure_aws(self, access_key: str, secret_key: str, region: str):
        """Configure AWS"""
        self.aws = AWSTool(access_key, secret_key, region)
        logger.info(f"AWS configured: {region}")

    def configure_jira(self, jira_url: str, email: str, api_token: str):
        """Configure Jira"""
        if UniversalJiraAutomation:
            self.jira = UniversalJiraAutomation(jira_url, email, api_token)
            logger.info("Jira configured")
        else:
            logger.warning("Jira tool not available")
        
        # Configure visual automation
        try:
            from .tools import JiraBrowserAutomation
            self.jira_visual = JiraBrowserAutomation(jira_url, email, api_token)
            logger.info("Jira visual automation configured")
        except ImportError:
            logger.warning("Jira visual automation not available")
    
    def configure_github(self, access_token: str):
        """Configure GitHub"""
        if GitHubTool:
            self.github = GitHubTool(access_token)
            logger.info("GitHub configured")
        else:
            logger.warning("GitHub tool not available")
    
    def configure_vercel(self, token: str = None):
        """Configure Vercel"""
        if VercelTool:
            self.vercel = VercelTool(token)
            logger.info("Vercel configured")
        else:
            logger.warning("Vercel tool not available")
    def configure_figma(self, token: str = None):
        """Configure Figma"""
        if FigmaToWebsiteTool:
            self.figma = FigmaToWebsiteTool(token)
            logger.info("Figma to Website tool configured")
        else:
            logger.warning("Figma tool not available")
    
    def configure_document_generator(self, ai_client=None):
        """Configure Document Generator"""
        if DocumentGeneratorTool:
            self.doc_gen = DocumentGeneratorTool(ai_client)
            logger.info("Document Generator configured")
        else:
            logger.warning("Document Generator not available")
    
    def _init_screenshot_doc(self):
        """Initialize Word document for screenshots"""
        if not self.screenshot_doc:
            Path(self.screenshot_dir).mkdir(exist_ok=True)
            self.screenshot_doc = docx.Document()
            self.screenshot_doc.add_heading('Automation Screenshots', 0)
            self.screenshot_doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            self.screenshot_doc.add_page_break()
    
    def _add_screenshot_to_doc(self, title: str, screenshot_path: str):
        """Add screenshot to Word document"""
        if self.screenshot_doc and Path(screenshot_path).exists():
            self.screenshot_doc.add_heading(title, level=1)
            self.screenshot_doc.add_paragraph(f'Time: {datetime.now().strftime("%H:%M:%S")}')
            try:
                self.screenshot_doc.add_picture(screenshot_path, width=Inches(6))
            except Exception as e:
                self.screenshot_doc.add_paragraph(f'[Image error: {e}]')
            self.screenshot_doc.add_page_break()
    
    def _save_screenshot_doc(self) -> str:
        """Save and return path to Word document"""
        if self.screenshot_doc:
            doc_path = f"automation_screenshots_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            self.screenshot_doc.save(doc_path)
            return doc_path
        return None
    
    def _register_tools(self):
        """Register MCP tools"""
        
        # Define all tools (including new GitHub and Vercel tools)
        self.tools_list = [
            # Browser tools
            Tool(
                name="browser_navigate",
                description="Navigate to a URL in the browser",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "url": {"type": "string", "description": "URL to navigate to"}
                    },
                    "required": ["url"]
                }
            ),
            Tool(
                name="youtube_play",
                description="Search and play a video on YouTube",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Video search query"}
                    },
                    "required": ["query"]
                }
            ),
            Tool(
                name="browser_screenshot",
                description="Take a screenshot of the current page",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "filename": {
                            "type": "string", 
                            "description": "Optional filename for screenshot"
                        }
                    }
                }
            ),
            
            # AWS tools
            Tool(
                name="aws_create_bucket",
                description="Create an AWS S3 bucket",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "bucket_name": {
                            "type": "string",
                            "description": "Name of the bucket to create"
                        }
                    },
                    "required": ["bucket_name"]
                }
            ),
            Tool(
                name="aws_list_buckets",
                description="List all AWS S3 buckets",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            
            # File system tools
            Tool(
                name="file_read",
                description="Read contents of a file",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to the file"}
                    },
                    "required": ["path"]
                }
            ),
            Tool(
                name="file_write",
                description="Write content to a file",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to the file"},
                        "content": {"type": "string", "description": "Content to write"}
                    },
                    "required": ["path", "content"]
                }
            ),
            Tool(
                name="list_directory",
                description="List contents of a directory",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {
                            "type": "string",
                            "description": "Directory path (default: current directory)"
                        }
                    }
                }
            ),
             Tool(
                name="figma_to_website",
                description="Convert Figma design to complete deployed website with all files saved locally, GitHub repo, and Vercel deployment",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "figma_url": {
                            "type": "string",
                            "description": "Figma file URL (e.g., https://www.figma.com/file/...)"
                        },
                        "project_name": {
                            "type": "string",
                            "description": "Name for the project"
                        },
                        "description": {
                            "type": "string",
                            "description": "Project description (optional)"
                        }
                    },
                    "required": ["figma_url", "project_name"]
                }
            ),
            
            # Document & Presentation Generator tools (NEW)
            Tool(
                name="create_document_and_presentation",
                description="Generate professional Word document and PowerPoint presentation from natural language request with AI research. Saves all files locally.",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "request": {
                            "type": "string",
                            "description": "What document to create (e.g., 'create project proposal for fitness app')"
                        },
                        "project_name": {
                            "type": "string",
                            "description": "Project name (optional, auto-generated from request if not provided)"
                        },
                        "output_folder": {
                            "type": "string",
                            "description": "Custom output folder path (optional, uses 'documents/' by default)"
                        }
                    },
                    "required": ["request"]
                }
            ),
            # Jira tools
            Tool(
                name="jira_create_assignment",
                description="Create the complete College Event Management Jira assignment with all epics, stories, and tasks",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            Tool(
                name="jira_create_projects",
                description="Create CEA and EAP projects in Jira",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            Tool(
                name="jira_create_visual_with_screenshots",
                description="Create Jira assignment VISUALLY with browser automation - takes screenshots at each step and creates Word document",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "document_path": {
                            "type": "string",
                            "description": "Path to assignment document (optional, uses default if not provided)"
                        }
                    }
                }
            ),
            Tool(
                name="jira_create_visual",
                description="Create Jira assignment VISUALLY using browser automation - watch it happen step by step",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "document_text": {
                            "type": "string",
                            "description": "Full text of the assignment document"
                        }
                    },
                    "required": ["document_text"]
                }
            ),
            
            # GitHub tools (NEW)
            Tool(
                name="github_create_repo",
                description="Create a new GitHub repository",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "repo_name": {
                            "type": "string",
                            "description": "Name of the repository to create"
                        },
                        "description": {
                            "type": "string",
                            "description": "Repository description (optional)"
                        },
                        "private": {
                            "type": "boolean",
                            "description": "Make repository private (default: false)"
                        }
                    },
                    "required": ["repo_name"]
                }
            ),
            Tool(
                name="github_push_code",
                description="Push local code to GitHub repository",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "repo_name": {
                            "type": "string",
                            "description": "Name of the repository"
                        },
                        "local_path": {
                            "type": "string",
                            "description": "Local path to push (default: current directory)"
                        },
                        "branch": {
                            "type": "string",
                            "description": "Branch name (default: main)"
                        }
                    },
                    "required": ["repo_name"]
                }
            ),
            Tool(
                name="github_generate_readme",
                description="Generate README.md file for the project based on structure",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "local_path": {
                            "type": "string",
                            "description": "Local path (default: current directory)"
                        }
                    }
                }
            ),
            Tool(
                name="github_build_and_push",
                description="Complete workflow: Generate README, create repo, and push code to GitHub",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "repo_name": {
                            "type": "string",
                            "description": "Name of the repository"
                        },
                        "description": {
                            "type": "string",
                            "description": "Repository description (optional)"
                        },
                        "local_path": {
                            "type": "string",
                            "description": "Local path (default: current directory)"
                        }
                    },
                    "required": ["repo_name"]
                }
            ),
            
            # Vercel tools (NEW)
            Tool(
                name="vercel_deploy",
                description="Deploy project to Vercel",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "local_path": {
                            "type": "string",
                            "description": "Local path to deploy (default: current directory)"
                        },
                        "production": {
                            "type": "boolean",
                            "description": "Deploy to production (default: true)"
                        }
                    }
                }
            ),
            Tool(
                name="vercel_list_deployments",
                description="List recent Vercel deployments",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "limit": {
                            "type": "integer",
                            "description": "Number of deployments to list (default: 5)"
                        }
                    }
                }
            ),
            
            # Screenshot tools
            Tool(
                name="enable_screenshots",
                description="Enable automatic screenshots for all subsequent actions",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            Tool(
                name="disable_screenshots",
                description="Disable automatic screenshots and save the Word document",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            Tool(
                name="take_screenshot_and_save",
                description="Take a screenshot of current browser page and add to Word document",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "title": {
                            "type": "string",
                            "description": "Title for this screenshot in the document"
                        }
                    },
                    "required": ["title"]
                }
            ),
        ]
        
        logger.info(f"Registered {len(self.tools_list)} MCP tools")
        
        # Register handlers with MCP server
        @self.server.list_tools()
        async def list_tools() -> list[Tool]:
            """Return list of available tools"""
            return self.tools_list
        
        @self.server.call_tool()
        async def call_tool(name: str, arguments: dict) -> list[TextContent]:
            """Execute a tool"""
            logger.info(f"Tool called: {name} with args: {arguments}")
            
            try:
                result = await self.execute_tool(name, arguments)
                
                # Format response
                if result.get("success"):
                    message = f"✅ {result.get('message', 'Success')}"
                    if result.get('data'):
                        data = result['data']
                        # Format different data types
                        if 'url' in data:
                            message += f"\n🔗 URL: {data['url']}"
                        if 'repo_url' in data:
                            message += f"\n📦 Repository: {data['repo_url']}"
                        if 'clone_url' in data:
                            message += f"\n📥 Clone: {data['clone_url']}"
                        if 'steps' in data:
                            message += f"\n📋 Steps:\n" + "\n".join(data['steps'])
                        if 'doc_path' in data:
                            message += f"\n📄 Document: {data['doc_path']}"
                else:
                    message = f"❌ Error: {result.get('error', 'Unknown error')}"
                
                return [TextContent(type="text", text=message)]
                
            except Exception as e:
                logger.error(f"Tool execution error: {e}", exc_info=True)
                return [TextContent(
                    type="text",
                    text=f"❌ Tool execution failed: {str(e)}"
                )]
    
    async def _auto_screenshot(self, title: str):
        """Auto-take screenshot if enabled"""
        if self.screenshot_enabled and self.browser.page:
            await self._take_and_save_screenshot(title)
    
    async def _take_and_save_screenshot(self, title: str) -> str:
        """Take screenshot and add to Word doc"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        screenshot_path = f"{self.screenshot_dir}/screenshot_{timestamp}.png"
        
        await self.browser.page.screenshot(path=screenshot_path)
        self._add_screenshot_to_doc(title, screenshot_path)
        
        return screenshot_path
    
    async def _jira_visual_with_screenshots(self, document_path: str = None):
        """Create Jira assignment visually with screenshots"""
        try:
            if not self.jira_visual:
                return {"success": False, "error": "Jira visual automation not configured"}
            
            # Parse document
            from .tools import DocumentParser
            parser = DocumentParser()
            
            if not document_path:
                document_path = 'documents/Assignment-11 Monday.docx'
            
            parse_result = await parser.parse_document(document_path, use_ai=False)
            
            if not parse_result['success']:
                return {"success": False, "error": f"Parse failed: {parse_result['error']}"}
            
            config = parse_result['data']
            
            # Initialize screenshot doc
            self._init_screenshot_doc()
            
            # Start browser
            await self.jira_visual.start_browser(headless=False)
            
            # Login
            await self.jira_visual.login_to_jira()
            screenshot_path = f"{self.screenshot_dir}/01_login.png"
            await self.jira_visual.page.screenshot(path=screenshot_path)
            self._add_screenshot_to_doc("1. Login to Jira", screenshot_path)
            await self.jira_visual.page.wait_for_timeout(2000)
            
            # Create Projects
            for idx, proj in enumerate(config.get('projects', []), start=1):
                await self.jira_visual.create_project_visual(proj['key'], proj['name'])
                screenshot_path = f"{self.screenshot_dir}/02_project_{proj['key']}.png"
                await self.jira_visual.page.screenshot(path=screenshot_path)
                self._add_screenshot_to_doc(f"2.{idx}. Project: {proj['key']}", screenshot_path)
                await self.jira_visual.page.wait_for_timeout(2000)
            
            # Create Epics
            for idx, epic in enumerate(config.get('epics', []), start=1):
                await self.jira_visual.create_epic_visual(epic['project'], epic['name'])
                screenshot_path = f"{self.screenshot_dir}/03_epic_{idx}.png"
                await self.jira_visual.page.screenshot(path=screenshot_path)
                self._add_screenshot_to_doc(f"3.{idx}. Epic: {epic['name']}", screenshot_path)
                await self.jira_visual.page.wait_for_timeout(1500)
            
            # Create Stories
            for idx, story in enumerate(config.get('stories', []), start=1):
                await self.jira_visual.create_story_visual(
                    story['project'],
                    story['summary'],
                    epic_name=story.get('epic_id'),
                    labels=story.get('labels'),
                    story_points=story.get('story_points')
                )
                screenshot_path = f"{self.screenshot_dir}/04_story_{idx}.png"
                await self.jira_visual.page.screenshot(path=screenshot_path)
                self._add_screenshot_to_doc(f"4.{idx}. Story: {story['summary'][:50]}", screenshot_path)
                await self.jira_visual.page.wait_for_timeout(1000)
            
            # Final screenshot
            screenshot_path = f"{self.screenshot_dir}/05_final.png"
            await self.jira_visual.page.screenshot(path=screenshot_path)
            self._add_screenshot_to_doc("5. Final - All Created", screenshot_path)
            
            # Save Word doc
            doc_path = self._save_screenshot_doc()
            
            # Keep browser open briefly
            await self.jira_visual.page.wait_for_timeout(30000)
            
            # Close browser
            if self.jira_visual.browser:
                await self.jira_visual.browser.close()
            if self.jira_visual.playwright:
                await self.jira_visual.playwright.stop()
            
            return {
                "success": True,
                "message": f"Visual automation complete with screenshots",
                "data": {
                    "doc_path": doc_path,
                    "projects": len(config.get('projects', [])),
                    "epics": len(config.get('epics', [])),
                    "stories": len(config.get('stories', []))
                }
            }
            
        except Exception as e:
            logger.error(f"Visual automation failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def get_tools_list(self):
        """Get list of available tools"""
        return self.tools_list
    
    async def execute_tool(self, name: str, arguments: dict):
        """Execute a tool directly"""
        logger.info(f"Executing tool: {name} with args: {arguments}")
        
        try:
            result = None
            
            # Browser tools
            if name == "browser_navigate":
                result = await self.browser.navigate(arguments["url"])
                if self.screenshot_enabled:
                    await self._auto_screenshot(f"Navigated to: {arguments['url']}")
                
            elif name == "youtube_play":
                result = await self.browser.play_youtube(arguments["query"])
                if self.screenshot_enabled:
                    await self._auto_screenshot(f"Playing: {arguments['query']}")
                
            elif name == "browser_screenshot":
                filename = arguments.get("filename", "screenshot.png")
                result = await self.browser.screenshot(filename)
            
            # AWS tools
            elif name == "aws_create_bucket":
                if not self.aws:
                    result = {"success": False, "error": "AWS not configured"}
                else:
                    result = await self.aws.create_bucket(arguments["bucket_name"])
                    
            elif name == "aws_list_buckets":
                if not self.aws:
                    result = {"success": False, "error": "AWS not configured"}
                else:
                    result = await self.aws.list_buckets()
            
            # System tools
            elif name == "file_read":
                result = await self.system.read_file(arguments["path"])
                
            elif name == "file_write":
                result = await self.system.write_file(
                    arguments["path"],
                    arguments["content"]
                )
                
            elif name == "list_directory":
                path = arguments.get("path", ".")
                result = await self.system.list_directory(path)
            
            # Jira tools
            elif name == "jira_create_assignment":
                if not self.jira:
                    result = {"success": False, "error": "Jira not configured"}
                else:
                    result = await self.jira.create_complete_assignment()
                    
            elif name == "jira_create_projects":
                if not self.jira:
                    result = {"success": False, "error": "Jira not configured"}
                else:
                    result = await self.jira.create_projects()
            
            elif name == "jira_create_visual_with_screenshots":
                result = await self._jira_visual_with_screenshots(arguments.get("document_path"))
            
            elif name == "jira_create_visual":
                if not self.jira_visual:
                    result = {"success": False, "error": "Jira visual automation not configured"}
                else:
                    from .tools import DocumentParser
                    parser = DocumentParser()
                    parse_result = parser.parse_with_rules(arguments["document_text"])
                    result = await self.jira_visual.create_assignment_visual(parse_result)
            
            # GitHub tools (NEW)
            elif name == "github_create_repo":
                if not self.github:
                    result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env"}
                else:
                    result = await self.github.create_repo(
                        arguments["repo_name"],
                        arguments.get("description", ""),
                        arguments.get("private", False)
                    )
            
            elif name == "github_push_code":
                if not self.github:
                    result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env"}
                else:
                    result = await self.github.push_local_code(
                        arguments["repo_name"],
                        arguments.get("local_path", "."),
                        arguments.get("branch", "main")
                    )
            
            elif name == "github_generate_readme":
                if not self.github:
                    result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env"}
                else:
                    result = await self.github.generate_readme(
                        arguments.get("local_path", ".")
                    )
            
            elif name == "github_build_and_push":
                if not self.github:
                    result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env"}
                else:
                    result = await self.github.build_and_push_project(
                        arguments["repo_name"],
                        arguments.get("description", "MCP Automation Project"),
                        arguments.get("local_path", ".")
                    )
            
            # Vercel tools (NEW)
            elif name == "vercel_deploy":
                if not self.vercel:
                    result = {"success": False, "error": "Vercel not configured"}
                else:
                    result = await self.vercel.deploy(
                        arguments.get("local_path", ".") if arguments else ".",
                        arguments.get("production", True) if arguments else True
                    )
            
            elif name == "vercel_list_deployments":
                if not self.vercel:
                    result = {"success": False, "error": "Vercel not configured"}
                else:
                    result = await self.vercel.get_deployments(
                        arguments.get("limit", 5) if arguments else 5
                    )
             # Figma to Website tools (NEW)
            elif name == "figma_to_website":
                if not self.figma:
                    result = {"success": False, "error": "Figma tool not configured"}
                else:
                    result = await self.figma.create_website_from_figma(
                        arguments["figma_url"],
                        arguments["project_name"],
                        arguments.get("description", "Website from Figma")
                    )
            
            # Document Generator tools (NEW)
            elif name == "create_document_and_presentation":
                if not self.doc_gen:
                    result = {"success": False, "error": "Document Generator not configured"}
                else:
                    result = await self.doc_gen.create_document_and_presentation(
                        arguments["request"],
                        arguments.get("project_name"),
                        arguments.get("output_folder")
                    )

            
            # Screenshot tools
            elif name == "enable_screenshots":
                self.screenshot_enabled = True
                self._init_screenshot_doc()
                result = {"success": True, "message": "Screenshots enabled"}
            
            elif name == "disable_screenshots":
                self.screenshot_enabled = False
                doc_path = self._save_screenshot_doc()
                result = {"success": True, "message": f"Document saved: {doc_path}", "data": {"doc_path": doc_path}}
            
            elif name == "take_screenshot_and_save":
                if not self.browser.page:
                    result = {"success": False, "error": "Browser not started"}
                else:
                    if not self.screenshot_doc:
                        self._init_screenshot_doc()
                    screenshot_path = await self._take_and_save_screenshot(arguments.get("title", "Screenshot"))
                    result = {"success": True, "message": f"Screenshot saved: {screenshot_path}"}
            
            else:
                result = {"success": False, "error": f"Unknown tool: {name}"}
            
            # Log usage
            self.db.log_tool_usage(name, result.get("success", False), 0)
            
            return result
            
        except Exception as e:
            logger.error(f"Tool execution error: {e}", exc_info=True)
            return {"success": False, "error": str(e)}