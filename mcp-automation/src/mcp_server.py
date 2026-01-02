# """MCP Server implementation - UPDATED with Workflow Engine and Complete Features + Code Analysis + Web Scraping"""
# from mcp.server import Server
# from mcp.types import Tool, TextContent

# # FIXED: Support both relative and absolute imports
# try:
#     from .tools import BrowserTool, AWSTool, SystemTool
# except ImportError:
#     from tools import BrowserTool, AWSTool, SystemTool

# try:
#     from .tools import UniversalJiraAutomation
# except ImportError:
#     UniversalJiraAutomation = None

# try:
#     from .tools import GitHubTool
# except ImportError:
#     GitHubTool = None

# try:
#     from .tools import VercelTool
# except ImportError:
#     VercelTool = None

# try:
#     from .tools import DocumentGeneratorTool
# except ImportError:
#     DocumentGeneratorTool = None

# # NEW: Code Analysis and Web Scraper Tools
# try:
#     from .tools.code_analysis_tool import CodeAnalysisTool
# except ImportError:
#     try:
#         from tools.code_analysis_tool import CodeAnalysisTool
#     except ImportError:
#         CodeAnalysisTool = None

# try:
#     from .tools.web_scraper_tool import WebScraperTool
# except ImportError:
#     try:
#         from tools.web_scraper_tool import WebScraperTool
#     except ImportError:
#         WebScraperTool = None

# # FIXED: DatabaseManager import
# try:
#     from .database import DatabaseManager
# except ImportError:
#     from database import DatabaseManager

# # FIXED: Workflow Engine import
# try:
#     from .workflow_engine import WorkflowEngine
# except ImportError:
#     try:
#         from workflow_engine import WorkflowEngine
#     except ImportError:
#         WorkflowEngine = None

# import logging
# from pathlib import Path
# from datetime import datetime
# import docx
# import asyncio
# from docx.shared import Inches
# import os
# import json
# import websockets

# logger = logging.getLogger(__name__)


# def sanitize_project_name(name: str) -> str:
#     """
#     Sanitize project name for GitHub and Vercel compatibility
    
#     Rules:
#     - Lowercase only
#     - Letters, digits, '.', '_', '-' allowed  
#     - No sequence '---' (triple dash)
#     - Max 100 characters
#     - Cannot start/end with dash
#     """
#     import re
#     import time
    
#     # Convert to lowercase
#     name = name.lower()
    
#     # Replace invalid characters with dash
#     name = re.sub(r'[^a-z0-9._-]', '-', name)
    
#     # Replace multiple consecutive dashes with double dash
#     name = re.sub(r'-{3,}', '--', name)
    
#     # Remove leading/trailing dashes
#     name = name.strip('-')
    
#     # Ensure not empty
#     if not name:
#         name = f"project-{int(time.time())}"
    
#     # Limit to 100 characters
#     if len(name) > 100:
#         name = name[:100].rstrip('-')
    
#     return name



# class MCPServer:
#     """MCP Server managing automation tools with Workflow Engine"""
    
#     def __init__(self):
#         self.server = Server("automation-server")
#         self.browser = BrowserTool()
#         self.aws = None
#         self.system = SystemTool()
#         self.jira = None
#         self.jira_visual = None  # Visual automation
#         self.github = None  # GitHub tool
#         self.vercel = None  # Vercel tool
#         self.doc_gen = None  # Document generator tool
#         self.workflow_engine = None  # Workflow engine
#         self.code_analysis = None  # NEW: Code analysis tool
#         self.web_scraper = None  # NEW: Web scraper tool
#         self.db = DatabaseManager()
        
#         # Screenshot management
#         self.screenshot_enabled = False
#         self.screenshot_doc = None
#         self.screenshot_dir = "screenshots"
        
#         # Store tools list manually
#         self.tools_list = []
        
#         self._register_tools()
        
#         # Initialize workflow engine AFTER tools are registered
#         if WorkflowEngine:
#             self.workflow_engine = WorkflowEngine(self)
#             logger.info("✅ Workflow engine initialized")
#         else:
#             logger.warning("⚠️  Workflow engine not available")
    
#     def configure_aws(self, access_key: str, secret_key: str, region: str):
#         """Configure AWS"""
#         self.aws = AWSTool(access_key, secret_key, region)
#         logger.info(f"AWS configured: {region}")

#     def configure_jira(self, jira_url: str, email: str, api_token: str):
#         """Configure Jira"""
#         if UniversalJiraAutomation:
#             self.jira = UniversalJiraAutomation(jira_url, email, api_token)
#             logger.info("Jira configured")
#         else:
#             logger.warning("Jira tool not available")
        
#         # Configure visual automation
#         try:
#             from .tools import JiraBrowserAutomation
#         except ImportError:
#             try:
#                 from tools import JiraBrowserAutomation
#             except ImportError:
#                 JiraBrowserAutomation = None
        
#         if JiraBrowserAutomation:
#             self.jira_visual = JiraBrowserAutomation(jira_url, email, api_token)
#             logger.info("Jira visual automation configured")
#         else:
#             logger.warning("Jira visual automation not available")
    
#     def configure_github(self, access_token: str):
#         """Configure GitHub"""
#         if GitHubTool:
#             self.github = GitHubTool(access_token)
#             logger.info("GitHub configured")
#         else:
#             logger.warning("GitHub tool not available")
    
#     def configure_vercel(self, token: str = None):
#         """Configure Vercel"""
#         if VercelTool:
#             self.vercel = VercelTool(token)
#             logger.info("Vercel configured")
#         else:
#             logger.warning("Vercel tool not available")
    
#     def configure_document_generator(self, ai_client=None):
#         """Configure Document Generator"""
#         if DocumentGeneratorTool:
#             self.doc_gen = DocumentGeneratorTool(ai_client)
#             logger.info("Document Generator configured")
#         else:
#             logger.warning("Document Generator not available")
    
#     def configure_code_analysis(self, ai_client):
#         """Configure Code Analysis Tool"""
#         if CodeAnalysisTool:
#             self.code_analysis = CodeAnalysisTool(ai_client)
#             logger.info("✅ Code Analysis Tool configured")
#         else:
#             logger.warning("⚠️  Code Analysis Tool not available")

#     def configure_web_scraper(self, ai_client):
#         """Configure Web Scraper Tool"""
#         if WebScraperTool:
#             self.web_scraper = WebScraperTool(ai_client)
#             logger.info("✅ Web Scraper Tool configured")
#         else:
#             logger.warning("⚠️  Web Scraper Tool not available")
    
#     def _init_screenshot_doc(self):
#         """Initialize Word document for screenshots"""
#         if not self.screenshot_doc:
#             Path(self.screenshot_dir).mkdir(exist_ok=True)
#             self.screenshot_doc = docx.Document()
#             self.screenshot_doc.add_heading('Automation Screenshots', 0)
#             self.screenshot_doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
#             self.screenshot_doc.add_page_break()
    
#     def _add_screenshot_to_doc(self, title: str, screenshot_path: str):
#         """Add screenshot to Word document"""
#         if self.screenshot_doc and Path(screenshot_path).exists():
#             self.screenshot_doc.add_heading(title, level=1)
#             self.screenshot_doc.add_paragraph(f'Time: {datetime.now().strftime("%H:%M:%S")}')
#             try:
#                 self.screenshot_doc.add_picture(screenshot_path, width=Inches(6))
#             except Exception as e:
#                 self.screenshot_doc.add_paragraph(f'[Image error: {e}]')
#             self.screenshot_doc.add_page_break()
    
#     def _save_screenshot_doc(self) -> str:
#         """Save and return path to Word document"""
#         if self.screenshot_doc:
#             doc_path = f"automation_screenshots_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
#             self.screenshot_doc.save(doc_path)
#             return doc_path
#         return None
    
#     def _register_tools(self):
#         """Register MCP tools"""
        
#         # Define all tools (including new workflow tools + code analysis + web scraping)
#         self.tools_list = [
#             # Browser tools
#             Tool(
#                 name="browser_navigate",
#                 description="Navigate to a URL in the browser",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "url": {"type": "string", "description": "URL to navigate to"}
#                     },
#                     "required": ["url"]
#                 }
#             ),
#             Tool(
#                 name="youtube_play",
#                 description="Search and play a video on YouTube",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "query": {"type": "string", "description": "Video search query"}
#                     },
#                     "required": ["query"]
#                 }
#             ),
#             Tool(
#                 name="browser_screenshot",
#                 description="Take a screenshot of the current page",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "filename": {
#                             "type": "string", 
#                             "description": "Optional filename for screenshot"
#                         }
#                     }
#                 }
#             ),
            
#             # AWS tools
#             Tool(
#                 name="aws_create_bucket",
#                 description="Create an AWS S3 bucket",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "bucket_name": {
#                             "type": "string",
#                             "description": "Name of the bucket to create"
#                         }
#                     },
#                     "required": ["bucket_name"]
#                 }
#             ),
#             Tool(
#                 name="aws_list_buckets",
#                 description="List all AWS S3 buckets",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {}
#                 }
#             ),
            
#             # File system tools
#             Tool(
#                 name="file_read",
#                 description="Read contents of a file",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "path": {"type": "string", "description": "Path to the file"}
#                     },
#                     "required": ["path"]
#                 }
#             ),
#             Tool(
#                 name="file_write",
#                 description="Write content to a file",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "path": {"type": "string", "description": "Path to the file"},
#                         "content": {"type": "string", "description": "Content to write"}
#                     },
#                     "required": ["path", "content"]
#                 }
#             ),
#             Tool(
#                 name="list_directory",
#                 description="List contents of a directory",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "path": {
#                             "type": "string",
#                             "description": "Directory path (default: current directory)"
#                         }
#                     }
#                 }
#             ),
            
#             # Figma to Website
#             Tool(
#                 name="figma_to_website",
#                 description="Convert Figma design to complete deployed website with all files saved locally, GitHub repo, and Vercel deployment",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "figma_url": {
#                             "type": "string",
#                             "description": "Figma file URL (e.g., https://www.figma.com/file/...)"
#                         },
#                         "project_name": {
#                             "type": "string",
#                             "description": "Name for the project"
#                         },
#                         "description": {
#                             "type": "string",
#                             "description": "Project description (optional)"
#                         }
#                     },
#                     "required": ["figma_url", "project_name"]
#                 }
#             ),
            
#             # Document & Presentation Generator
#             Tool(
#                 name="create_document_and_presentation",
#                 description="Generate professional Word document and PowerPoint presentation from natural language request with AI research. Saves all files locally.",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "request": {
#                             "type": "string",
#                             "description": "What document to create (e.g., 'create project proposal for fitness app')"
#                         },
#                         "project_name": {
#                             "type": "string",
#                             "description": "Project name (optional, auto-generated from request if not provided)"
#                         },
#                         "output_folder": {
#                             "type": "string",
#                             "description": "Custom output folder path (optional, uses 'documents/' by default)"
#                         }
#                     },
#                     "required": ["request"]
#                 }
#             ),
            
#             # Jira tools
#             Tool(
#                 name="jira_create_assignment",
#                 description="Create the complete College Event Management Jira assignment with all epics, stories, and tasks",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {}
#                 }
#             ),
#             Tool(
#                 name="jira_create_projects",
#                 description="Create CEA and EAP projects in Jira",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {}
#                 }
#             ),
#             Tool(
#                 name="jira_create_visual_with_screenshots",
#                 description="Create Jira assignment VISUALLY with browser automation - takes screenshots at each step and creates Word document",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "document_path": {
#                             "type": "string",
#                             "description": "Path to assignment document (optional, uses default if not provided)"
#                         }
#                     }
#                 }
#             ),
#             Tool(
#                 name="jira_create_visual",
#                 description="Create Jira assignment VISUALLY using browser automation - watch it happen step by step",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "document_text": {
#                             "type": "string",
#                             "description": "Full text of the assignment document"
#                         }
#                     },
#                     "required": ["document_text"]
#                 }
#             ),
            
#             # GitHub tools
#             Tool(
#                 name="github_create_repo",
#                 description="Create a new GitHub repository",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "repo_name": {
#                             "type": "string",
#                             "description": "Name of the repository to create"
#                         },
#                         "description": {
#                             "type": "string",
#                             "description": "Repository description (optional)"
#                         },
#                         "private": {
#                             "type": "boolean",
#                             "description": "Make repository private (default: false)"
#                         }
#                     },
#                     "required": ["repo_name"]
#                 }
#             ),
#             Tool(
#                 name="github_push_code",
#                 description="Push local code to GitHub repository",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "repo_name": {
#                             "type": "string",
#                             "description": "Name of the repository"
#                         },
#                         "local_path": {
#                             "type": "string",
#                             "description": "Local path to push (default: current directory)"
#                         },
#                         "branch": {
#                             "type": "string",
#                             "description": "Branch name (default: main)"
#                         }
#                     },
#                     "required": ["repo_name"]
#                 }
#             ),
#             Tool(
#                 name="github_generate_readme",
#                 description="Generate README.md file for the project based on structure",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "project_path": {
#                             "type": "string",
#                             "description": "Project path"
#                         },
#                         "project_name": {
#                             "type": "string",
#                             "description": "Project name"
#                         }
#                     },
#                     "required": ["project_path", "project_name"]
#                 }
#             ),
#             Tool(
#                 name="github_build_and_push",
#                 description="Complete workflow: Generate README, create repo, and push code to GitHub",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "repo_name": {
#                             "type": "string",
#                             "description": "Name of the repository"
#                         },
#                         "description": {
#                             "type": "string",
#                             "description": "Repository description (optional)"
#                         },
#                         "local_path": {
#                             "type": "string",
#                             "description": "Local path (default: current directory)"
#                         }
#                     },
#                     "required": ["repo_name"]
#                 }
#             ),
            
#             # Vercel tools
#             Tool(
#                 name="vercel_deploy",
#                 description="Deploy project to Vercel",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "local_path": {
#                             "type": "string",
#                             "description": "Local path to deploy (default: current directory)"
#                         },
#                         "production": {
#                             "type": "boolean",
#                             "description": "Deploy to production (default: true)"
#                         }
#                     }
#                 }
#             ),
#             Tool(
#                 name="vercel_list_deployments",
#                 description="List recent Vercel deployments",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "limit": {
#                             "type": "integer",
#                             "description": "Number of deployments to list (default: 5)"
#                         }
#                     }
#                 }
#             ),

#             # Add these at the end of self.tools_list (before the closing bracket)

# # ⭐ NEW: Windows System Automation Tools
# Tool(
#     name="transfer_files",
#     description="Transfer/copy files from one location to another (C: to D:, Downloads to Desktop, etc). Works with selected files in File Explorer or specified paths.",
#     inputSchema={
#         "type": "object",
#         "properties": {
#             "source": {"type": "string", "description": "Source folder path (e.g., C:\\Users\\tarun\\Downloads)"},
#             "destination": {"type": "string", "description": "Destination folder path (e.g., D:\\Backup)"},
#             "files": {
#                 "type": "array", 
#                 "items": {"type": "string"}, 
#                 "description": "Specific files to transfer (optional, transfers all if not specified)"
#             }
#         },
#         "required": ["source", "destination"]
#     }
# ),
# Tool(
#     name="move_files",
#     description="Move files from one location to another (removes from source)",
#     inputSchema={
#         "type": "object",
#         "properties": {
#             "source": {"type": "string", "description": "Source folder path"},
#             "destination": {"type": "string", "description": "Destination folder path"},
#             "files": {
#                 "type": "array",
#                 "items": {"type": "string"},
#                 "description": "Specific files to move (optional)"
#             }
#         },
#         "required": ["source", "destination"]
#     }
# ),
# Tool(
#     name="install_app",
#     description="Install an application from Microsoft Store",
#     inputSchema={
#         "type": "object",
#         "properties": {
#             "app_name": {"type": "string", "description": "Name of the app to install (e.g., 'WhatsApp', 'Spotify')"}
#         },
#         "required": ["app_name"]
#     }
# ),
# Tool(
#     name="launch_app",
#     description="Launch a Windows application",
#     inputSchema={
#         "type": "object",
#         "properties": {
#             "app_name": {
#                 "type": "string", 
#                 "description": "Name of the app to launch (notepad, calculator, chrome, vscode, etc)"
#             }
#         },
#         "required": ["app_name"]
#     }
# ),
# Tool(
#     name="close_app",
#     description="Close a running Windows application",
#     inputSchema={
#         "type": "object",
#         "properties": {
#             "app_name": {"type": "string", "description": "Name of the app to close"}
#         },
#         "required": ["app_name"]
#     }
# ),
# Tool(
#     name="get_system_info",
#     description="Get Windows system information (CPU, RAM, disk space, running processes)",
#     inputSchema={
#         "type": "object",
#         "properties": {},
#         "required": []
#     }
# ),
# Tool(
#     name="get_running_apps",
#     description="Get list of all currently running applications",
#     inputSchema={
#         "type": "object",
#         "properties": {},
#         "required": []
#     }
# ),
# Tool(
#     name="open_folder",
#     description="Open a folder in Windows File Explorer",
#     inputSchema={
#         "type": "object",
#         "properties": {
#             "path": {"type": "string", "description": "Folder path to open"}
#         },
#         "required": ["path"]
#     }
# ),
# Tool(
#     name="create_folder",
#     description="Create a new folder at specified path",
#     inputSchema={
#         "type": "object",
#         "properties": {
#             "path": {"type": "string", "description": "Full path for new folder"}
#         },
#         "required": ["path"]
#     }
# ),
            
#             # ⭐ NEW: Workflow Engine Tools
#             Tool(
#                 name="execute_workflow",
#                 description="Execute a multi-step workflow. Available workflows: figma_to_production (Figma→Code→GitHub→Vercel), document_code (Code→Documentation), setup_project (GitHub repo setup)",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "workflow_name": {
#                             "type": "string",
#                             "description": "Name of workflow: figma_to_production, document_code, or setup_project",
#                             "enum": ["figma_to_production", "document_code", "setup_project"]
#                         },
#                         "parameters": {
#                             "type": "object",
#                             "description": "Workflow parameters (varies by workflow)"
#                         }
#                     },
#                     "required": ["workflow_name", "parameters"]
#                 }
#             ),
#             Tool(
#                 name="list_workflows",
#                 description="List all available multi-step workflows with descriptions",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {}
#                 }
#             ),
            
#             # Screenshot tools
#             Tool(
#                 name="enable_screenshots",
#                 description="Enable automatic screenshots for all subsequent actions",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {}
#                 }
#             ),
#             Tool(
#                 name="disable_screenshots",
#                 description="Disable automatic screenshots and save the Word document",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {}
#                 }
#             ),
#             Tool(
#                 name="take_screenshot_and_save",
#                 description="Take a screenshot of current browser page and add to Word document",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "title": {
#                             "type": "string",
#                             "description": "Title for this screenshot in the document"
#                         }
#                     },
#                     "required": ["title"]
#                 }
#             ),
            
#             # ⭐ NEW: Code Analysis Tools
#             Tool(
#                 name="analyze_code",
#                 description="Analyze code file for bugs, quality issues, and improvements. Use with current file path from context.",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "file_path": {
#                             "type": "string",
#                             "description": "Path to code file to analyze"
#                         }
#                     },
#                     "required": ["file_path"]
#                 }
#             ),
#             Tool(
#                 name="fix_bugs",
#                 description="Suggest bug fixes for code in current file",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "file_path": {
#                             "type": "string",
#                             "description": "Path to code file"
#                         },
#                         "context": {
#                             "type": "string",
#                             "description": "Additional context about the bug"
#                         }
#                     },
#                     "required": ["file_path"]
#                 }
#             ),
#             Tool(
#                 name="generate_tests",
#                 description="Generate unit tests for code in current file",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "file_path": {
#                             "type": "string",
#                             "description": "Path to code file"
#                         }
#                     },
#                     "required": ["file_path"]
#                 }
#             ),
#             Tool(
#                 name="refactor_code",
#                 description="Refactor code for better quality and maintainability",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "file_path": {
#                             "type": "string",
#                             "description": "Path to code file"
#                         }
#                     },
#                     "required": ["file_path"]
#                 }
#             ),
#             Tool(
#                 name="document_code",
#                 description="Generate comprehensive documentation for code in current file with docstrings",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "file_path": {
#                             "type": "string",
#                             "description": "Path to code file"
#                         }
#                     },
#                     "required": ["file_path"]
#                 }
#             ),
            
#             # ⭐ NEW: Web Scraping Tools
#             Tool(
#                 name="summarize_article",
#                 description="Extract and summarize article from current webpage or URL using AI",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "url": {
#                             "type": "string",
#                             "description": "URL of article (optional if already on page)"
#                         }
#                     }
#                 }
#             ),
#             Tool(
#                 name="scrape_table_to_csv",
#                 description="Scrape HTML table from current page and save to CSV file",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "selector": {
#                             "type": "string",
#                             "description": "CSS selector for table (optional, finds first table if not provided)"
#                         },
#                         "filename": {
#                             "type": "string",
#                             "description": "Output CSV filename (default: scraped_table.csv)"
#                         }
#                     }
#                 }
#             ),
#             Tool(
#                 name="research_competitors",
#                 description="Research competitors by searching Google and analyzing results with AI",
#                 inputSchema={
#                     "type": "object",
#                     "properties": {
#                         "topic": {
#                             "type": "string",
#                             "description": "Topic to research (e.g., 'login page design', 'e-commerce checkout')"
#                         },
#                         "num_results": {
#                             "type": "integer",
#                             "description": "Number of competitor results to analyze (default: 5, max: 10)"
#                         }
#                     },
#                     "required": ["topic"]
#                 }
#             ),
#         ]
        
#         logger.info(f"Registered {len(self.tools_list)} MCP tools")
        
#         # Register handlers with MCP server
#         @self.server.list_tools()
#         async def list_tools() -> list[Tool]:
#             """Return list of available tools"""
#             return self.tools_list
        
#         @self.server.call_tool()
#         async def call_tool(name: str, arguments: dict) -> list[TextContent]:
#             """Execute a tool"""
#             logger.info(f"Tool called: {name} with args: {arguments}")
            
#             try:
#                 result = await self.execute_tool(name, arguments)
                
#                 # Format response
#                 if result.get("success"):
#                     message = f"✅ {result.get('message', 'Success')}"
#                     if result.get('data'):
#                         data = result['data']
#                         # Format different data types
#                         if 'url' in data:
#                             message += f"\n🔗 URL: {data['url']}"
#                         if 'repo_url' in data:
#                             message += f"\n📦 Repository: {data['repo_url']}"
#                         if 'clone_url' in data:
#                             message += f"\n📥 Clone: {data['clone_url']}"
#                         if 'steps' in data:
#                             message += f"\n📋 Steps:\n" + "\n".join(data['steps'])
#                         if 'doc_path' in data:
#                             message += f"\n📄 Document: {data['doc_path']}"
#                         if 'workflows' in data:
#                             message += f"\n📋 Available Workflows:\n"
#                             for wf in data['workflows']:
#                                 message += f"  • {wf['name']}: {wf['description']} ({wf['steps']} steps)\n"
#                 else:
#                     message = f"❌ Error: {result.get('error', 'Unknown error')}"
                
#                 return [TextContent(type="text", text=message)]
                
#             except Exception as e:
#                 logger.error(f"Tool execution error: {e}", exc_info=True)
#                 return [TextContent(
#                     type="text",
#                     text=f"❌ Tool execution failed: {str(e)}"
#                 )]
    
#     async def _auto_screenshot(self, title: str):
#         """Auto-take screenshot if enabled"""
#         if self.screenshot_enabled and self.browser.page:
#             await self._take_and_save_screenshot(title)
    
#     async def _take_and_save_screenshot(self, title: str) -> str:
#         """Take screenshot and add to Word doc"""
#         timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
#         screenshot_path = f"{self.screenshot_dir}/screenshot_{timestamp}.png"
        
#         await self.browser.page.screenshot(path=screenshot_path)
#         self._add_screenshot_to_doc(title, screenshot_path)
        
#         return screenshot_path
    
#     async def _jira_visual_with_screenshots(self, document_path: str = None):
#         """Create Jira assignment visually with screenshots"""
#         try:
#             if not self.jira_visual:
#                 return {"success": False, "error": "Jira visual automation not configured"}
            
#             # Parse document
#             try:
#                 from .tools import DocumentParser
#             except ImportError:
#                 from tools import DocumentParser
            
#             parser = DocumentParser()
            
#             if not document_path:
#                 document_path = 'documents/Assignment-11 Monday.docx'
            
#             parse_result = await parser.parse_document(document_path, use_ai=False)
            
#             if not parse_result['success']:
#                 return {"success": False, "error": f"Parse failed: {parse_result['error']}"}
            
#             config = parse_result['data']
            
#             # Initialize screenshot doc
#             self._init_screenshot_doc()
            
#             # Start browser
#             await self.jira_visual.start_browser(headless=False)
            
#             # Login
#             await self.jira_visual.login_to_jira()
#             screenshot_path = f"{self.screenshot_dir}/01_login.png"
#             await self.jira_visual.page.screenshot(path=screenshot_path)
#             self._add_screenshot_to_doc("1. Login to Jira", screenshot_path)
#             await self.jira_visual.page.wait_for_timeout(2000)
            
#             # Create Projects
#             for idx, proj in enumerate(config.get('projects', []), start=1):
#                 await self.jira_visual.create_project_visual(proj['key'], proj['name'])
#                 screenshot_path = f"{self.screenshot_dir}/02_project_{proj['key']}.png"
#                 await self.jira_visual.page.screenshot(path=screenshot_path)
#                 self._add_screenshot_to_doc(f"2.{idx}. Project: {proj['key']}", screenshot_path)
#                 await self.jira_visual.page.wait_for_timeout(2000)
            
#             # Create Epics
#             for idx, epic in enumerate(config.get('epics', []), start=1):
#                 await self.jira_visual.create_epic_visual(epic['project'], epic['name'])
#                 screenshot_path = f"{self.screenshot_dir}/03_epic_{idx}.png"
#                 await self.jira_visual.page.screenshot(path=screenshot_path)
#                 self._add_screenshot_to_doc(f"3.{idx}. Epic: {epic['name']}", screenshot_path)
#                 await self.jira_visual.page.wait_for_timeout(1500)
            
#             # Create Stories
#             for idx, story in enumerate(config.get('stories', []), start=1):
#                 await self.jira_visual.create_story_visual(
#                     story['project'],
#                     story['summary'],
#                     epic_name=story.get('epic_id'),
#                     labels=story.get('labels'),
#                     story_points=story.get('story_points')
#                 )
#                 screenshot_path = f"{self.screenshot_dir}/04_story_{idx}.png"
#                 await self.jira_visual.page.screenshot(path=screenshot_path)
#                 self._add_screenshot_to_doc(f"4.{idx}. Story: {story['summary'][:50]}", screenshot_path)
#                 await self.jira_visual.page.wait_for_timeout(1000)
            
#             # Final screenshot
#             screenshot_path = f"{self.screenshot_dir}/05_final.png"
#             await self.jira_visual.page.screenshot(path=screenshot_path)
#             self._add_screenshot_to_doc("5. Final - All Created", screenshot_path)
            
#             # Save Word doc
#             doc_path = self._save_screenshot_doc()
            
#             # Keep browser open briefly
#             await self.jira_visual.page.wait_for_timeout(30000)
            
#             # Close browser
#             if self.jira_visual.browser:
#                 await self.jira_visual.browser.close()
#             if self.jira_visual.playwright:
#                 await self.jira_visual.playwright.stop()
            
#             return {
#                 "success": True,
#                 "completed": True,
#                 "message": f"Visual automation complete with screenshots",
#                 "data": {
#                     "doc_path": doc_path,
#                     "projects": len(config.get('projects', [])),
#                     "epics": len(config.get('epics', [])),
#                     "stories": len(config.get('stories', []))
#                 }
#             }
            
#         except Exception as e:
#             logger.error(f"Visual automation failed: {e}")
#             return {"success": False, "error": str(e), "fatal": True}
    
#     async def get_tools_list(self):
#         """Get list of available tools"""
#         return self.tools_list
    
#     async def execute_tool(self, name: str, arguments: dict):
#         """Execute a tool directly"""
#         logger.info(f"Executing tool: {name} with args: {arguments}")
        
#         try:
#             result = None
            
#             # Browser tools
#             if name == "browser_navigate":
#                 result = await self.browser.navigate(arguments["url"])
#                 if self.screenshot_enabled:
#                     await self._auto_screenshot(f"Navigated to: {arguments['url']}")
                
#             elif name == "youtube_play":
#                 result = await self.browser.play_youtube(arguments["query"])
#                 if self.screenshot_enabled:
#                     await self._auto_screenshot(f"Playing: {arguments['query']}")
                
#             elif name == "browser_screenshot":
#                 filename = arguments.get("filename", "screenshot.png")
#                 result = await self.browser.screenshot(filename)
            
#             # AWS tools
#             elif name == "aws_create_bucket":
#                 if not self.aws:
#                     result = {"success": False, "error": "AWS not configured", "fatal": True}
#                 else:
#                     result = await self.aws.create_bucket(arguments["bucket_name"])
                    
#             elif name == "aws_list_buckets":
#                 if not self.aws:
#                     result = {"success": False, "error": "AWS not configured", "fatal": True}
#                 else:
#                     result = await self.aws.list_buckets()
            
#             # System tools
#             elif name == "file_read":
#                 result = await self.system.read_file(arguments["path"])
                
#             elif name == "file_write":
#                 result = await self.system.write_file(
#                     arguments["path"],
#                     arguments["content"]
#                 )
                
#             elif name == "list_directory":
#                 path = arguments.get("path", ".")
#                 result = await self.system.list_directory(path)
            
#             # Jira tools
#             elif name == "jira_create_assignment":
#                 if not self.jira:
#                     result = {"success": False, "error": "Jira not configured", "fatal": True}
#                 else:
#                     result = await self.jira.create_complete_assignment()
                    
#             elif name == "jira_create_projects":
#                 if not self.jira:
#                     result = {"success": False, "error": "Jira not configured", "fatal": True}
#                 else:
#                     result = await self.jira.create_projects()
            
#             elif name == "jira_create_visual_with_screenshots":
#                 result = await self._jira_visual_with_screenshots(arguments.get("document_path"))
            
#             elif name == "jira_create_visual":
#                 if not self.jira_visual:
#                     result = {"success": False, "error": "Jira visual automation not configured", "fatal": True}
#                 else:
#                     try:
#                         from .tools import DocumentParser
#                     except ImportError:
#                         from tools import DocumentParser
                    
#                     parser = DocumentParser()
#                     parse_result = parser.parse_with_rules(arguments["document_text"])
#                     result = await self.jira_visual.create_assignment_visual(parse_result)
            
#             # GitHub tools
#             elif name == "github_create_repo":
#                 if not self.github:
#                     result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env", "fatal": True}
#                 else:
#                     result = await self.github.create_repo(
#                         arguments["repo_name"],
#                         arguments.get("description", ""),
#                         arguments.get("private", False)
#                     )
            
#             elif name == "github_push_code":
#                 if not self.github:
#                     result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env", "fatal": True}
#                 else:
#                     result = await self.github.push_local_code(
#                         arguments["repo_name"],
#                         arguments.get("local_path", "."),
#                         arguments.get("branch", "main")
#                     )
            
#             elif name == "github_generate_readme":
#                 if not self.github:
#                     result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env", "fatal": True}
#                 else:
#                     result = await self.github.generate_readme(
#                         arguments["project_path"],
#                         arguments["project_name"]
#                     )
            
#             elif name == "github_build_and_push":
#                 if not self.github:
#                     result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env", "fatal": True}
#                 else:
#                     result = await self.github.build_and_push_project(
#                         arguments["repo_name"],
#                         arguments.get("description", "MCP Automation Project"),
#                         arguments.get("local_path", ".")
#                     )
            
#             # Vercel tools
#             elif name == "vercel_deploy":
#                 if not self.vercel:
#                     result = {"success": False, "error": "Vercel not configured", "fatal": True}
#                 else:
#                     result = await self.vercel.deploy(
#                         arguments.get("local_path", ".") if arguments else ".",
#                         arguments.get("production", True) if arguments else True
#                     )
            
#             elif name == "vercel_list_deployments":
#                 if not self.vercel:
#                     result = {"success": False, "error": "Vercel not configured", "fatal": True}
#                 else:
#                     result = await self.vercel.get_deployments(
#                         arguments.get("limit", 5) if arguments else 5
#                     )
            
#             # Workflow Engine Tools
#             elif name == "execute_workflow":
#                 if not self.workflow_engine:
#                     result = {"success": False, "error": "Workflow engine not initialized", "fatal": True}
#                 else:
#                     workflow_name = arguments["workflow_name"]
#                     params = arguments.get("parameters", {})
                    
#                     logger.info(f"🔄 Executing workflow: {workflow_name}")
#                     result = await self.workflow_engine.execute_workflow(workflow_name, params)
            
#             elif name == "list_workflows":
#                 if not self.workflow_engine:
#                     result = {"success": False, "error": "Workflow engine not initialized", "fatal": True}
#                 else:
#                     workflows = self.workflow_engine.get_available_workflows()
#                     result = {
#                         "success": True,
#                         "completed": True,
#                         "message": f"Found {len(workflows)} available workflows",
#                         "data": {"workflows": workflows}
#                     }
            
#             # Figma to Website
#             elif name == "figma_to_website":
#                 figma_token = os.getenv("FIGMA_TOKEN")
#                 if not figma_token:
#                     result = {
#                         "success": False, 
#                         "error": "FIGMA_TOKEN not set in .env file", 
#                         "fatal": True
#                     }
#                 else:
#                     try:
#                         try:
#                             from .tools.production_figma_converter import ProductionFigmaToCode
#                         except ImportError:
#                             from tools.production_figma_converter import ProductionFigmaToCode
                        
#                         figma_url = arguments["figma_url"]
#                         project_name_raw = arguments["project_name"]
                        
#                         # Sanitize project name
#                         project_name = sanitize_project_name(project_name_raw)
#                         logger.info(f"📝 Original name: '{project_name_raw}' → Sanitized: '{project_name}'")
                        
#                         output_path = Path("projects") / project_name
                        
#                         logger.info(f"🎨 Step 1/3: Converting Figma design to code...")
                        
#                         # STEP 1: Convert Figma to Code
#                         converter = ProductionFigmaToCode(figma_token)
#                         conversion_result = await converter.convert(figma_url, output_path)
                        
#                         if not conversion_result.get("success"):
#                             result = {
#                                 "success": False,
#                                 "error": f"Figma conversion failed: {conversion_result.get('error')}",
#                                 "fatal": True
#                             }
#                         else:
#                             steps = ["✅ Converted Figma design to code"]
#                             data = {
#                                 "local_path": str(output_path),
#                                 "project_name": project_name
#                             }
                            
#                             # STEP 2: Push to GitHub
#                             if self.github:
#                                 logger.info(f"📦 Step 2/3: Pushing to GitHub...")
#                                 try:
#                                     create_result = await self.github.create_repo(
#                                         project_name,
#                                         description=f"Website generated from Figma design",
#                                         private=False
#                                     )
                                    
#                                     if create_result.get("success"):
#                                         push_result = await self.github.push_local_code(
#                                             project_name,
#                                             str(output_path),
#                                             branch="main"
#                                         )
                                        
#                                         if push_result.get("success"):
#                                             steps.append(f"✅ Pushed to GitHub")
#                                             data["github_url"] = push_result.get("data", {}).get("repo_url")
#                                             data["repo_name"] = project_name
#                                         else:
#                                             steps.append(f"⚠️ GitHub push failed: {push_result.get('error')}")
#                                     else:
#                                         steps.append(f"⚠️ GitHub repo creation failed: {create_result.get('error')}")
                                        
#                                 except Exception as e:
#                                     steps.append(f"⚠️ GitHub error: {str(e)}")
#                                     logger.error(f"GitHub error: {e}", exc_info=True)
#                             else:
#                                 steps.append("⚠️ Skipped GitHub (GITHUB_TOKEN not set)")
                            
#                             # STEP 3: Deploy to Vercel
#                             if self.vercel:
#                                 logger.info(f"🚀 Step 3/3: Deploying to Vercel...")
#                                 try:
#                                     deploy_result = await self.vercel.deploy(
#                                         str(output_path),
#                                         production=True
#                                     )
                                    
#                                     if deploy_result.get("success"):
#                                         vercel_url = deploy_result.get("data", {}).get("url")
#                                         steps.append(f"✅ Deployed to Vercel")
#                                         data["vercel_url"] = vercel_url
#                                         data["deployment_url"] = vercel_url
#                                     else:
#                                         steps.append(f"⚠️ Vercel deployment failed: {deploy_result.get('error')}")
                                        
#                                 except Exception as e:
#                                     steps.append(f"⚠️ Vercel error: {str(e)}")
#                                     logger.error(f"Vercel error: {e}", exc_info=True)
#                             else:
#                                 steps.append("⚠️ Skipped Vercel (VERCEL_TOKEN not set)")
                            
#                             # Success!
#                             result = {
#                                 "success": True,
#                                 "completed": True,
#                                 "message": "Website created successfully!",
#                                 "steps": steps,
#                                 "data": data
#                             }
                            
#                             logger.info(f"🎉 Complete workflow finished!")
                        
#                     except ImportError as e:
#                         logger.error(f"Import error: {e}", exc_info=True)
#                         result = {
#                             "success": False,
#                             "error": f"Failed to import Figma converter: {str(e)}",
#                             "fatal": True
#                         }
#                     except Exception as e:
#                         logger.error(f"Figma workflow error: {e}", exc_info=True)
#                         result = {
#                             "success": False,
#                             "error": f"Workflow failed: {str(e)}",
#                             "fatal": True
#                         }
            
#             # Document Generator
#             elif name == "create_document_and_presentation":
#                 if not self.doc_gen:
#                     result = {"success": False, "error": "Document Generator not configured", "fatal": True}
#                 else:
#                     result = await self.doc_gen.create_document_and_presentation(
#                         arguments["request"],
#                         arguments.get("project_name"),
#                         arguments.get("output_folder")
#                     )
            
#             # Screenshot tools
#             elif name == "enable_screenshots":
#                 self.screenshot_enabled = True
#                 self._init_screenshot_doc()
#                 result = {"success": True, "completed": True, "message": "Screenshots enabled"}
            
#             elif name == "disable_screenshots":
#                 self.screenshot_enabled = False
#                 doc_path = self._save_screenshot_doc()
#                 result = {"success": True, "completed": True, "message": f"Document saved: {doc_path}", "data": {"doc_path": doc_path}}
            
#             elif name == "take_screenshot_and_save":
#                 if not self.browser.page:
#                     result = {"success": False, "error": "Browser not started"}
#                 else:
#                     if not self.screenshot_doc:
#                         self._init_screenshot_doc()
#                     screenshot_path = await self._take_and_save_screenshot(arguments.get("title", "Screenshot"))
#                     result = {"success": True, "completed": True, "message": f"Screenshot saved: {screenshot_path}"}
            
#             # ⭐ NEW: Code Analysis tools
#             elif name == "analyze_code":
#                 if not self.code_analysis:
#                     result = {"success": False, "error": "Code analysis not configured. AI client required.", "fatal": True}
#                 else:
#                     result = await self.code_analysis.analyze_code(arguments["file_path"])

#             elif name == "fix_bugs":
#                 if not self.code_analysis:
#                     result = {"success": False, "error": "Code analysis not configured. AI client required.", "fatal": True}
#                 else:
#                     result = await self.code_analysis.fix_bugs(
#                         arguments["file_path"],
#                         arguments.get("context")
#                     )

#             elif name == "generate_tests":
#                 if not self.code_analysis:
#                     result = {"success": False, "error": "Code analysis not configured. AI client required.", "fatal": True}
#                 else:
#                     result = await self.code_analysis.generate_tests(arguments["file_path"])

#             elif name == "refactor_code":
#                 if not self.code_analysis:
#                     result = {"success": False, "error": "Code analysis not configured. AI client required.", "fatal": True}
#                 else:
#                     result = await self.code_analysis.refactor_code(arguments["file_path"])

#             elif name == "document_code":
#                 if not self.code_analysis:
#                     result = {"success": False, "error": "Code analysis not configured. AI client required.", "fatal": True}
#                 else:
#                     result = await self.code_analysis.document_code(arguments["file_path"])

#             # ⭐ NEW: Web Scraping tools
#             elif name == "summarize_article":
#                 if not self.web_scraper:
#                     result = {"success": False, "error": "Web scraper not configured. AI client required.", "fatal": True}
#                 else:
#                     result = await self.web_scraper.summarize_article(arguments.get("url"))

#             elif name == "scrape_table_to_csv":
#                 if not self.web_scraper:
#                     result = {"success": False, "error": "Web scraper not configured.", "fatal": True}
#                 else:
#                     result = await self.web_scraper.scrape_table_to_csv(
#                         arguments.get("selector"),
#                         arguments.get("filename", "scraped_table.csv")
#                     )

#             elif name == "research_competitors":
#                 if not self.web_scraper:
#                     result = {"success": False, "error": "Web scraper not configured. AI client required.", "fatal": True}
#                 else:
#                     result = await self.web_scraper.research_competitors(
#                         arguments["topic"],
#                         arguments.get("num_results", 5)
#                     )
            
#             else:
#                 result = {"success": False, "error": f"Unknown tool: {name}", "fatal": True}


                


            

            







            
#             # Log usage
#             self.db.log_tool_usage(name, result.get("success", False), 0)
            
#             # CHECK FOR COMPLETION OR FATAL ERROR
#             if result.get("completed") is True:
#                 logger.info(f"✅ Tool '{name}' marked as completed")
#             elif result.get("fatal") is True:
#                 logger.error(f"❌ Tool '{name}' returned fatal error")
            
#             return result
            
#         except Exception as e:
#             logger.error(f"Tool execution error: {e}", exc_info=True)
#             return {"success": False, "error": str(e), "fatal": True}



"""MCP Server implementation - UPDATED with Workflow Engine and Complete Features + Code Analysis + Web Scraping"""
from mcp.server import Server
from mcp.types import Tool, TextContent

# FIXED: Support both relative and absolute imports
try:
    from .tools import BrowserTool, AWSTool, SystemTool
except ImportError:
    from tools import BrowserTool, AWSTool, SystemTool

try:
    from .tools import UniversalJiraAutomation
except ImportError:
    UniversalJiraAutomation = None

try:
    from .tools import GitHubTool
except ImportError:
    GitHubTool = None

try:
    from .tools import VercelTool
except ImportError:
    VercelTool = None

try:
    from .tools import DocumentGeneratorTool
except ImportError:
    DocumentGeneratorTool = None

# NEW: Code Analysis and Web Scraper Tools
try:
    from .tools.code_analysis_tool import CodeAnalysisTool
except ImportError:
    try:
        from tools.code_analysis_tool import CodeAnalysisTool
    except ImportError:
        CodeAnalysisTool = None

try:
    from .tools.web_scraper_tool import WebScraperTool
except ImportError:
    try:
        from tools.web_scraper_tool import WebScraperTool
    except ImportError:
        WebScraperTool = None

# FIXED: DatabaseManager import
try:
    from .database import DatabaseManager
except ImportError:
    from database import DatabaseManager

# FIXED: Workflow Engine import
try:
    from .workflow_engine import WorkflowEngine
except ImportError:
    try:
        from workflow_engine import WorkflowEngine
    except ImportError:
        WorkflowEngine = None

import logging
from pathlib import Path
from datetime import datetime
import docx
import asyncio
from docx.shared import Inches
import os
import json
import websockets

logger = logging.getLogger(__name__)


def sanitize_project_name(name: str) -> str:
    """
    Sanitize project name for GitHub and Vercel compatibility
    
    Rules:
    - Lowercase only
    - Letters, digits, '.', '_', '-' allowed  
    - No sequence '---' (triple dash)
    - Max 100 characters
    - Cannot start/end with dash
    """
    import re
    import time
    
    # Convert to lowercase
    name = name.lower()
    
    # Replace invalid characters with dash
    name = re.sub(r'[^a-z0-9._-]', '-', name)
    
    # Replace multiple consecutive dashes with double dash
    name = re.sub(r'-{3,}', '--', name)
    
    # Remove leading/trailing dashes
    name = name.strip('-')
    
    # Ensure not empty
    if not name:
        name = f"project-{int(time.time())}"
    
    # Limit to 100 characters
    if len(name) > 100:
        name = name[:100].rstrip('-')
    
    return name



class MCPServer:
    """MCP Server managing automation tools with Workflow Engine"""
    
    def __init__(self):
        self.server = Server("automation-server")
        self.browser = BrowserTool()
        self.aws = None
        self.system = SystemTool()
        self.jira = None
        self.jira_visual = None  # Visual automation
        self.github = None  # GitHub tool
        self.vercel = None  # Vercel tool
        self.doc_gen = None  # Document generator tool
        self.workflow_engine = None  # Workflow engine
        self.code_analysis = None  # NEW: Code analysis tool
        self.web_scraper = None  # NEW: Web scraper tool
        self.db = DatabaseManager()
        
        # Screenshot management
        self.screenshot_enabled = False
        self.screenshot_doc = None
        self.screenshot_dir = "screenshots"
        
        # Store tools list manually
        self.tools_list = []
        
        self._register_tools()
        
        # Initialize workflow engine AFTER tools are registered
        if WorkflowEngine:
            self.workflow_engine = WorkflowEngine(self)
            logger.info("✅ Workflow engine initialized")
        else:
            logger.warning("⚠️  Workflow engine not available")
    
    def configure_aws(self, access_key: str, secret_key: str, region: str):
        """Configure AWS"""
        self.aws = AWSTool(access_key, secret_key, region)
        logger.info(f"AWS configured: {region}")

    def configure_jira(self, jira_url: str, email: str, api_token: str):
        """Configure Jira"""
        if UniversalJiraAutomation:
            self.jira = UniversalJiraAutomation(jira_url, email, api_token)
            logger.info("Jira configured")
        else:
            logger.warning("Jira tool not available")
        
        # Configure visual automation
        try:
            from .tools import JiraBrowserAutomation
        except ImportError:
            try:
                from tools import JiraBrowserAutomation
            except ImportError:
                JiraBrowserAutomation = None
        
        if JiraBrowserAutomation:
            self.jira_visual = JiraBrowserAutomation(jira_url, email, api_token)
            logger.info("Jira visual automation configured")
        else:
            logger.warning("Jira visual automation not available")
    
    def configure_github(self, access_token: str):
        """Configure GitHub"""
        if GitHubTool:
            self.github = GitHubTool(access_token)
            logger.info("GitHub configured")
        else:
            logger.warning("GitHub tool not available")
    
    def configure_vercel(self, token: str = None):
        """Configure Vercel"""
        if VercelTool:
            self.vercel = VercelTool(token)
            logger.info("Vercel configured")
        else:
            logger.warning("Vercel tool not available")
    
    def configure_document_generator(self, ai_client=None):
        """Configure Document Generator"""
        if DocumentGeneratorTool:
            self.doc_gen = DocumentGeneratorTool(ai_client)
            logger.info("Document Generator configured")
        else:
            logger.warning("Document Generator not available")
    
    def configure_code_analysis(self, ai_client):
        """Configure Code Analysis Tool"""
        if CodeAnalysisTool:
            self.code_analysis = CodeAnalysisTool(ai_client)
            logger.info("✅ Code Analysis Tool configured")
        else:
            logger.warning("⚠️  Code Analysis Tool not available")

    def configure_web_scraper(self, ai_client):
        """Configure Web Scraper Tool"""
        if WebScraperTool:
            self.web_scraper = WebScraperTool(ai_client)
            logger.info("✅ Web Scraper Tool configured")
        else:
            logger.warning("⚠️  Web Scraper Tool not available")
    
    def _init_screenshot_doc(self):
        """Initialize Word document for screenshots"""
        if not self.screenshot_doc:
            Path(self.screenshot_dir).mkdir(exist_ok=True)
            self.screenshot_doc = docx.Document()
            self.screenshot_doc.add_heading('Automation Screenshots', 0)
            self.screenshot_doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            self.screenshot_doc.add_page_break()
    
    def _add_screenshot_to_doc(self, title: str, screenshot_path: str):
        """Add screenshot to Word document"""
        if self.screenshot_doc and Path(screenshot_path).exists():
            self.screenshot_doc.add_heading(title, level=1)
            self.screenshot_doc.add_paragraph(f'Time: {datetime.now().strftime("%H:%M:%S")}')
            try:
                self.screenshot_doc.add_picture(screenshot_path, width=Inches(6))
            except Exception as e:
                self.screenshot_doc.add_paragraph(f'[Image error: {e}]')
            self.screenshot_doc.add_page_break()
    
    def _save_screenshot_doc(self) -> str:
        """Save and return path to Word document"""
        if self.screenshot_doc:
            doc_path = f"automation_screenshots_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            self.screenshot_doc.save(doc_path)
            return doc_path
        return None
    
    def _register_tools(self):
        """Register MCP tools"""
        
        # Define all tools (including new workflow tools + code analysis + web scraping)
        self.tools_list = [
            # Browser tools
            Tool(
                name="browser_navigate",
                description="Navigate to a URL in the browser",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "url": {"type": "string", "description": "URL to navigate to"}
                    },
                    "required": ["url"]
                }
            ),
            Tool(
                name="youtube_play",
                description="Search and play a video on YouTube",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Video search query"}
                    },
                    "required": ["query"]
                }
            ),
            Tool(
                name="browser_screenshot",
                description="Take a screenshot of the current page",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "filename": {
                            "type": "string", 
                            "description": "Optional filename for screenshot"
                        }
                    }
                }
            ),
            
            # AWS tools
            Tool(
                name="aws_create_bucket",
                description="Create an AWS S3 bucket",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "bucket_name": {
                            "type": "string",
                            "description": "Name of the bucket to create"
                        }
                    },
                    "required": ["bucket_name"]
                }
            ),
            Tool(
                name="aws_list_buckets",
                description="List all AWS S3 buckets",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            
            # File system tools
            Tool(
                name="file_read",
                description="Read contents of a file",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to the file"}
                    },
                    "required": ["path"]
                }
            ),
            Tool(
                name="file_write",
                description="Write content to a file",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to the file"},
                        "content": {"type": "string", "description": "Content to write"}
                    },
                    "required": ["path", "content"]
                }
            ),
            Tool(
                name="list_directory",
                description="List contents of a directory",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {
                            "type": "string",
                            "description": "Directory path (default: current directory)"
                        }
                    }
                }
            ),
            
            # Figma to Website
            Tool(
                name="figma_to_website",
                description="Convert Figma design to complete deployed website with all files saved locally, GitHub repo, and Vercel deployment",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "figma_url": {
                            "type": "string",
                            "description": "Figma file URL (e.g., https://www.figma.com/file/...)"
                        },
                        "project_name": {
                            "type": "string",
                            "description": "Name for the project"
                        },
                        "description": {
                            "type": "string",
                            "description": "Project description (optional)"
                        }
                    },
                    "required": ["figma_url", "project_name"]
                }
            ),
            
            # Document & Presentation Generator
            Tool(
                name="create_document_and_presentation",
                description="Generate professional Word document and PowerPoint presentation from natural language request with AI research. Saves all files locally.",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "request": {
                            "type": "string",
                            "description": "What document to create (e.g., 'create project proposal for fitness app')"
                        },
                        "project_name": {
                            "type": "string",
                            "description": "Project name (optional, auto-generated from request if not provided)"
                        },
                        "output_folder": {
                            "type": "string",
                            "description": "Custom output folder path (optional, uses 'documents/' by default)"
                        }
                    },
                    "required": ["request"]
                }
            ),
            
            # Jira tools
            Tool(
                name="jira_create_assignment",
                description="Create the complete College Event Management Jira assignment with all epics, stories, and tasks",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            Tool(
                name="jira_create_projects",
                description="Create CEA and EAP projects in Jira",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            Tool(
                name="jira_create_visual_with_screenshots",
                description="Create Jira assignment VISUALLY with browser automation - takes screenshots at each step and creates Word document",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "document_path": {
                            "type": "string",
                            "description": "Path to assignment document (optional, uses default if not provided)"
                        }
                    }
                }
            ),
            Tool(
                name="jira_create_visual",
                description="Create Jira assignment VISUALLY using browser automation - watch it happen step by step",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "document_text": {
                            "type": "string",
                            "description": "Full text of the assignment document"
                        }
                    },
                    "required": ["document_text"]
                }
            ),
            
            # GitHub tools
            Tool(
                name="github_create_repo",
                description="Create a new GitHub repository",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "repo_name": {
                            "type": "string",
                            "description": "Name of the repository to create"
                        },
                        "description": {
                            "type": "string",
                            "description": "Repository description (optional)"
                        },
                        "private": {
                            "type": "boolean",
                            "description": "Make repository private (default: false)"
                        }
                    },
                    "required": ["repo_name"]
                }
            ),
            Tool(
                name="github_push_code",
                description="Push local code to GitHub repository",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "repo_name": {
                            "type": "string",
                            "description": "Name of the repository"
                        },
                        "local_path": {
                            "type": "string",
                            "description": "Local path to push (default: current directory)"
                        },
                        "branch": {
                            "type": "string",
                            "description": "Branch name (default: main)"
                        }
                    },
                    "required": ["repo_name"]
                }
            ),
            Tool(
                name="github_generate_readme",
                description="Generate README.md file for the project based on structure",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "project_path": {
                            "type": "string",
                            "description": "Project path"
                        },
                        "project_name": {
                            "type": "string",
                            "description": "Project name"
                        }
                    },
                    "required": ["project_path", "project_name"]
                }
            ),
            Tool(
                name="github_build_and_push",
                description="Complete workflow: Generate README, create repo, and push code to GitHub",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "repo_name": {
                            "type": "string",
                            "description": "Name of the repository"
                        },
                        "description": {
                            "type": "string",
                            "description": "Repository description (optional)"
                        },
                        "local_path": {
                            "type": "string",
                            "description": "Local path (default: current directory)"
                        }
                    },
                    "required": ["repo_name"]
                }
            ),
            
            # Vercel tools
            Tool(
                name="vercel_deploy",
                description="Deploy project to Vercel",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "local_path": {
                            "type": "string",
                            "description": "Local path to deploy (default: current directory)"
                        },
                        "production": {
                            "type": "boolean",
                            "description": "Deploy to production (default: true)"
                        }
                    }
                }
            ),
            Tool(
                name="vercel_list_deployments",
                description="List recent Vercel deployments",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "limit": {
                            "type": "integer",
                            "description": "Number of deployments to list (default: 5)"
                        }
                    }
                }
            ),

            # Add these at the end of self.tools_list (before the closing bracket)

# ⭐ NEW: Windows System Automation Tools
Tool(
    name="transfer_files",
    description="Transfer/copy files from one location to another (C: to D:, Downloads to Desktop, etc). Works with selected files in File Explorer or specified paths.",
    inputSchema={
        "type": "object",
        "properties": {
            "source": {"type": "string", "description": "Source folder path (e.g., C:\\Users\\tarun\\Downloads)"},
            "destination": {"type": "string", "description": "Destination folder path (e.g., D:\\Backup)"},
            "files": {
                "type": "array", 
                "items": {"type": "string"}, 
                "description": "Specific files to transfer (optional, transfers all if not specified)"
            }
        },
        "required": ["source", "destination"]
    }
),
Tool(
    name="move_files",
    description="Move files from one location to another (removes from source)",
    inputSchema={
        "type": "object",
        "properties": {
            "source": {"type": "string", "description": "Source folder path"},
            "destination": {"type": "string", "description": "Destination folder path"},
            "files": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Specific files to move (optional)"
            }
        },
        "required": ["source", "destination"]
    }
),
Tool(
    name="install_app",
    description="Install an application from Microsoft Store",
    inputSchema={
        "type": "object",
        "properties": {
            "app_name": {"type": "string", "description": "Name of the app to install (e.g., 'WhatsApp', 'Spotify')"}
        },
        "required": ["app_name"]
    }
),
Tool(
    name="launch_app",
    description="Launch a Windows application",
    inputSchema={
        "type": "object",
        "properties": {
            "app_name": {
                "type": "string", 
                "description": "Name of the app to launch (notepad, calculator, chrome, vscode, etc)"
            }
        },
        "required": ["app_name"]
    }
),
Tool(
    name="close_app",
    description="Close a running Windows application",
    inputSchema={
        "type": "object",
        "properties": {
            "app_name": {"type": "string", "description": "Name of the app to close"}
        },
        "required": ["app_name"]
    }
),
Tool(
    name="get_system_info",
    description="Get Windows system information (CPU, RAM, disk space, running processes)",
    inputSchema={
        "type": "object",
        "properties": {},
        "required": []
    }
),
Tool(
    name="get_running_apps",
    description="Get list of all currently running applications",
    inputSchema={
        "type": "object",
        "properties": {},
        "required": []
    }
),
Tool(
    name="open_folder",
    description="Open a folder in Windows File Explorer",
    inputSchema={
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Folder path to open"}
        },
        "required": ["path"]
    }
),
Tool(
    name="create_folder",
    description="Create a new folder at specified path",
    inputSchema={
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Full path for new folder"}
        },
        "required": ["path"]
    }
),
            
            # ⭐ NEW: Workflow Engine Tools
            Tool(
                name="execute_workflow",
                description="Execute a multi-step workflow. Available workflows: figma_to_production (Figma→Code→GitHub→Vercel), document_code (Code→Documentation), setup_project (GitHub repo setup)",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "workflow_name": {
                            "type": "string",
                            "description": "Name of workflow: figma_to_production, document_code, or setup_project",
                            "enum": ["figma_to_production", "document_code", "setup_project"]
                        },
                        "parameters": {
                            "type": "object",
                            "description": "Workflow parameters (varies by workflow)"
                        }
                    },
                    "required": ["workflow_name", "parameters"]
                }
            ),
            Tool(
                name="list_workflows",
                description="List all available multi-step workflows with descriptions",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            
            # Screenshot tools
            Tool(
                name="enable_screenshots",
                description="Enable automatic screenshots for all subsequent actions",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            Tool(
                name="disable_screenshots",
                description="Disable automatic screenshots and save the Word document",
                inputSchema={
                    "type": "object",
                    "properties": {}
                }
            ),
            Tool(
                name="take_screenshot_and_save",
                description="Take a screenshot of current browser page and add to Word document",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "title": {
                            "type": "string",
                            "description": "Title for this screenshot in the document"
                        }
                    },
                    "required": ["title"]
                }
            ),
            
            # ⭐ NEW: Code Analysis Tools
            Tool(
                name="analyze_code",
                description="Analyze code file for bugs, quality issues, and improvements. Use with current file path from context.",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "Path to code file to analyze"
                        }
                    },
                    "required": ["file_path"]
                }
            ),
            Tool(
                name="fix_bugs",
                description="Suggest bug fixes for code in current file",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "Path to code file"
                        },
                        "context": {
                            "type": "string",
                            "description": "Additional context about the bug"
                        }
                    },
                    "required": ["file_path"]
                }
            ),
            Tool(
                name="generate_tests",
                description="Generate unit tests for code in current file",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "Path to code file"
                        }
                    },
                    "required": ["file_path"]
                }
            ),
            Tool(
                name="refactor_code",
                description="Refactor code for better quality and maintainability",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "Path to code file"
                        }
                    },
                    "required": ["file_path"]
                }
            ),
            Tool(
                name="document_code",
                description="Generate comprehensive documentation for code in current file with docstrings",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "Path to code file"
                        }
                    },
                    "required": ["file_path"]
                }
            ),
            
            # ⭐ NEW: Web Scraping Tools
            Tool(
                name="summarize_article",
                description="Extract and summarize article from current webpage or URL using AI",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "url": {
                            "type": "string",
                            "description": "URL of article (optional if already on page)"
                        }
                    }
                }
            ),
            Tool(
                name="scrape_table_to_csv",
                description="Scrape HTML table from current page and save to CSV file",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "selector": {
                            "type": "string",
                            "description": "CSS selector for table (optional, finds first table if not provided)"
                        },
                        "filename": {
                            "type": "string",
                            "description": "Output CSV filename (default: scraped_table.csv)"
                        }
                    }
                }
            ),
            Tool(
                name="research_competitors",
                description="Research competitors by searching Google and analyzing results with AI",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "topic": {
                            "type": "string",
                            "description": "Topic to research (e.g., 'login page design', 'e-commerce checkout')"
                        },
                        "num_results": {
                            "type": "integer",
                            "description": "Number of competitor results to analyze (default: 5, max: 10)"
                        }
                    },
                    "required": ["topic"]
                }
            ),
        ]
        
        logger.info(f"Registered {len(self.tools_list)} MCP tools")
        
        # Register handlers with MCP server
        @self.server.list_tools()
        async def list_tools() -> list[Tool]:
            """Return list of available tools"""
            return self.tools_list
        
        @self.server.call_tool()
        async def call_tool(name: str, arguments: dict) -> list[TextContent]:
            """Execute a tool"""
            logger.info(f"Tool called: {name} with args: {arguments}")
            
            try:
                result = await self.execute_tool(name, arguments)
                
                # Format response
                if result.get("success"):
                    message = f"✅ {result.get('message', 'Success')}"
                    if result.get('data'):
                        data = result['data']
                        # Format different data types
                        if 'url' in data:
                            message += f"\n🔗 URL: {data['url']}"
                        if 'repo_url' in data:
                            message += f"\n📦 Repository: {data['repo_url']}"
                        if 'clone_url' in data:
                            message += f"\n📥 Clone: {data['clone_url']}"
                        if 'steps' in data:
                            message += f"\n📋 Steps:\n" + "\n".join(data['steps'])
                        if 'doc_path' in data:
                            message += f"\n📄 Document: {data['doc_path']}"
                        if 'workflows' in data:
                            message += f"\n📋 Available Workflows:\n"
                            for wf in data['workflows']:
                                message += f"  • {wf['name']}: {wf['description']} ({wf['steps']} steps)\n"
                else:
                    message = f"❌ Error: {result.get('error', 'Unknown error')}"
                
                return [TextContent(type="text", text=message)]
                
            except Exception as e:
                logger.error(f"Tool execution error: {e}", exc_info=True)
                return [TextContent(
                    type="text",
                    text=f"❌ Tool execution failed: {str(e)}"
                )]
    
    async def _auto_screenshot(self, title: str):
        """Auto-take screenshot if enabled"""
        if self.screenshot_enabled and self.browser.page:
            await self._take_and_save_screenshot(title)
    
    async def _take_and_save_screenshot(self, title: str) -> str:
        """Take screenshot and add to Word doc"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        screenshot_path = f"{self.screenshot_dir}/screenshot_{timestamp}.png"
        
        await self.browser.page.screenshot(path=screenshot_path)
        self._add_screenshot_to_doc(title, screenshot_path)
        
        return screenshot_path
    
    async def _jira_visual_with_screenshots(self, document_path: str = None):
        """Create Jira assignment visually with screenshots"""
        try:
            if not self.jira_visual:
                return {"success": False, "error": "Jira visual automation not configured"}
            
            # Parse document
            try:
                from .tools import DocumentParser
            except ImportError:
                from tools import DocumentParser
            
            parser = DocumentParser()
            
            if not document_path:
                document_path = 'documents/Assignment-11 Monday.docx'
            
            parse_result = await parser.parse_document(document_path, use_ai=False)
            
            if not parse_result['success']:
                return {"success": False, "error": f"Parse failed: {parse_result['error']}"}
            
            config = parse_result['data']
            
            # Initialize screenshot doc
            self._init_screenshot_doc()
            
            # Start browser
            await self.jira_visual.start_browser(headless=False)
            
            # Login
            await self.jira_visual.login_to_jira()
            screenshot_path = f"{self.screenshot_dir}/01_login.png"
            await self.jira_visual.page.screenshot(path=screenshot_path)
            self._add_screenshot_to_doc("1. Login to Jira", screenshot_path)
            await self.jira_visual.page.wait_for_timeout(2000)
            
            # Create Projects
            for idx, proj in enumerate(config.get('projects', []), start=1):
                await self.jira_visual.create_project_visual(proj['key'], proj['name'])
                screenshot_path = f"{self.screenshot_dir}/02_project_{proj['key']}.png"
                await self.jira_visual.page.screenshot(path=screenshot_path)
                self._add_screenshot_to_doc(f"2.{idx}. Project: {proj['key']}", screenshot_path)
                await self.jira_visual.page.wait_for_timeout(2000)
            
            # Create Epics
            for idx, epic in enumerate(config.get('epics', []), start=1):
                await self.jira_visual.create_epic_visual(epic['project'], epic['name'])
                screenshot_path = f"{self.screenshot_dir}/03_epic_{idx}.png"
                await self.jira_visual.page.screenshot(path=screenshot_path)
                self._add_screenshot_to_doc(f"3.{idx}. Epic: {epic['name']}", screenshot_path)
                await self.jira_visual.page.wait_for_timeout(1500)
            
            # Create Stories
            for idx, story in enumerate(config.get('stories', []), start=1):
                await self.jira_visual.create_story_visual(
                    story['project'],
                    story['summary'],
                    epic_name=story.get('epic_id'),
                    labels=story.get('labels'),
                    story_points=story.get('story_points')
                )
                screenshot_path = f"{self.screenshot_dir}/04_story_{idx}.png"
                await self.jira_visual.page.screenshot(path=screenshot_path)
                self._add_screenshot_to_doc(f"4.{idx}. Story: {story['summary'][:50]}", screenshot_path)
                await self.jira_visual.page.wait_for_timeout(1000)
            
            # Final screenshot
            screenshot_path = f"{self.screenshot_dir}/05_final.png"
            await self.jira_visual.page.screenshot(path=screenshot_path)
            self._add_screenshot_to_doc("5. Final - All Created", screenshot_path)
            
            # Save Word doc
            doc_path = self._save_screenshot_doc()
            
            # Keep browser open briefly
            await self.jira_visual.page.wait_for_timeout(30000)
            
            # Close browser
            if self.jira_visual.browser:
                await self.jira_visual.browser.close()
            if self.jira_visual.playwright:
                await self.jira_visual.playwright.stop()
            
            return {
                "success": True,
                "completed": True,
                "message": f"Visual automation complete with screenshots",
                "data": {
                    "doc_path": doc_path,
                    "projects": len(config.get('projects', [])),
                    "epics": len(config.get('epics', [])),
                    "stories": len(config.get('stories', []))
                }
            }
            
        except Exception as e:
            logger.error(f"Visual automation failed: {e}")
            return {"success": False, "error": str(e), "fatal": True}
    
    async def get_tools_list(self):
        """Get list of available tools"""
        return self.tools_list
    
    async def execute_tool(self, name: str, arguments: dict):
        """Execute a tool directly"""
        logger.info(f"Executing tool: {name} with args: {arguments}")
        
        try:
            result = None
            
            # Browser tools
            if name == "browser_navigate":
                result = await self.browser.navigate(arguments["url"])
                if self.screenshot_enabled:
                    await self._auto_screenshot(f"Navigated to: {arguments['url']}")
                
            elif name == "youtube_play":
                result = await self.browser.play_youtube(arguments["query"])
                if self.screenshot_enabled:
                    await self._auto_screenshot(f"Playing: {arguments['query']}")
                
            elif name == "browser_screenshot":
                filename = arguments.get("filename", "screenshot.png")
                result = await self.browser.screenshot(filename)
            
            # AWS tools
            elif name == "aws_create_bucket":
                if not self.aws:
                    result = {"success": False, "error": "AWS not configured", "fatal": True}
                else:
                    result = await self.aws.create_bucket(arguments["bucket_name"])
                    
            elif name == "aws_list_buckets":
                if not self.aws:
                    result = {"success": False, "error": "AWS not configured", "fatal": True}
                else:
                    result = await self.aws.list_buckets()
            
            # System tools
            elif name == "file_read":
                result = await self.system.read_file(arguments["path"])
                
            elif name == "file_write":
                result = await self.system.write_file(
                    arguments["path"],
                    arguments["content"]
                )
                
            elif name == "list_directory":
                path = arguments.get("path", ".")
                result = await self.system.list_directory(path)
            
            # Jira tools
            elif name == "jira_create_assignment":
                if not self.jira:
                    result = {"success": False, "error": "Jira not configured", "fatal": True}
                else:
                    result = await self.jira.create_complete_assignment()
                    
            elif name == "jira_create_projects":
                if not self.jira:
                    result = {"success": False, "error": "Jira not configured", "fatal": True}
                else:
                    result = await self.jira.create_projects()
            
            elif name == "jira_create_visual_with_screenshots":
                result = await self._jira_visual_with_screenshots(arguments.get("document_path"))
            
            elif name == "jira_create_visual":
                if not self.jira_visual:
                    result = {"success": False, "error": "Jira visual automation not configured", "fatal": True}
                else:
                    try:
                        from .tools import DocumentParser
                    except ImportError:
                        from tools import DocumentParser
                    
                    parser = DocumentParser()
                    parse_result = parser.parse_with_rules(arguments["document_text"])
                    result = await self.jira_visual.create_assignment_visual(parse_result)
            
            # GitHub tools
            elif name == "github_create_repo":
                if not self.github:
                    result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env", "fatal": True}
                else:
                    result = await self.github.create_repo(
                        arguments["repo_name"],
                        arguments.get("description", ""),
                        arguments.get("private", False)
                    )
            
            elif name == "github_push_code":
                if not self.github:
                    result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env", "fatal": True}
                else:
                    result = await self.github.push_local_code(
                        arguments["repo_name"],
                        arguments.get("local_path", "."),
                        arguments.get("branch", "main")
                    )
            
            elif name == "github_generate_readme":
                if not self.github:
                    result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env", "fatal": True}
                else:
                    result = await self.github.generate_readme(
                        arguments["project_path"],
                        arguments["project_name"]
                    )
            
            elif name == "github_build_and_push":
                if not self.github:
                    result = {"success": False, "error": "GitHub not configured. Add GITHUB_TOKEN to .env", "fatal": True}
                else:
                    result = await self.github.build_and_push_project(
                        arguments["repo_name"],
                        arguments.get("description", "MCP Automation Project"),
                        arguments.get("local_path", ".")
                    )
            
            # Vercel tools
            elif name == "vercel_deploy":
                if not self.vercel:
                    result = {"success": False, "error": "Vercel not configured", "fatal": True}
                else:
                    result = await self.vercel.deploy(
                        arguments.get("local_path", ".") if arguments else ".",
                        arguments.get("production", True) if arguments else True
                    )
            
            elif name == "vercel_list_deployments":
                if not self.vercel:
                    result = {"success": False, "error": "Vercel not configured", "fatal": True}
                else:
                    result = await self.vercel.get_deployments(
                        arguments.get("limit", 5) if arguments else 5
                    )
            
            # Workflow Engine Tools
            elif name == "execute_workflow":
                if not self.workflow_engine:
                    result = {"success": False, "error": "Workflow engine not initialized", "fatal": True}
                else:
                    workflow_name = arguments["workflow_name"]
                    params = arguments.get("parameters", {})
                    
                    logger.info(f"🔄 Executing workflow: {workflow_name}")
                    result = await self.workflow_engine.execute_workflow(workflow_name, params)
            
            elif name == "list_workflows":
                if not self.workflow_engine:
                    result = {"success": False, "error": "Workflow engine not initialized", "fatal": True}
                else:
                    workflows = self.workflow_engine.get_available_workflows()
                    result = {
                        "success": True,
                        "completed": True,
                        "message": f"Found {len(workflows)} available workflows",
                        "data": {"workflows": workflows}
                    }
            
            # Figma to Website
            elif name == "figma_to_website":
                figma_token = os.getenv("FIGMA_TOKEN")
                if not figma_token:
                    result = {
                        "success": False, 
                        "error": "FIGMA_TOKEN not set in .env file", 
                        "fatal": True
                    }
                else:
                    try:
                        try:
                            from .tools.production_figma_converter import ProductionFigmaToCode
                        except ImportError:
                            from tools.production_figma_converter import ProductionFigmaToCode
                        
                        figma_url = arguments["figma_url"]
                        project_name_raw = arguments["project_name"]
                        
                        # Sanitize project name
                        project_name = sanitize_project_name(project_name_raw)
                        logger.info(f"📝 Original name: '{project_name_raw}' → Sanitized: '{project_name}'")
                        
                        output_path = Path("projects") / project_name
                        
                        logger.info(f"🎨 Step 1/3: Converting Figma design to code...")
                        
                        # STEP 1: Convert Figma to Code
                        converter = ProductionFigmaToCode(figma_token)
                        conversion_result = await converter.convert(figma_url, output_path)
                        
                        if not conversion_result.get("success"):
                            result = {
                                "success": False,
                                "error": f"Figma conversion failed: {conversion_result.get('error')}",
                                "fatal": True
                            }
                        else:
                            steps = ["✅ Converted Figma design to code"]
                            data = {
                                "local_path": str(output_path),
                                "project_name": project_name
                            }
                            
                            # STEP 2: Push to GitHub
                            if self.github:
                                logger.info(f"📦 Step 2/3: Pushing to GitHub...")
                                try:
                                    create_result = await self.github.create_repo(
                                        project_name,
                                        description=f"Website generated from Figma design",
                                        private=False
                                    )
                                    
                                    if create_result.get("success"):
                                        push_result = await self.github.push_local_code(
                                            project_name,
                                            str(output_path),
                                            branch="main"
                                        )
                                        
                                        if push_result.get("success"):
                                            steps.append(f"✅ Pushed to GitHub")
                                            data["github_url"] = push_result.get("data", {}).get("repo_url")
                                            data["repo_name"] = project_name
                                        else:
                                            steps.append(f"⚠️ GitHub push failed: {push_result.get('error')}")
                                    else:
                                        steps.append(f"⚠️ GitHub repo creation failed: {create_result.get('error')}")
                                        
                                except Exception as e:
                                    steps.append(f"⚠️ GitHub error: {str(e)}")
                                    logger.error(f"GitHub error: {e}", exc_info=True)
                            else:
                                steps.append("⚠️ Skipped GitHub (GITHUB_TOKEN not set)")
                            
                            # STEP 3: Deploy to Vercel
                            if self.vercel:
                                logger.info(f"🚀 Step 3/3: Deploying to Vercel...")
                                try:
                                    deploy_result = await self.vercel.deploy(
                                        str(output_path),
                                        production=True
                                    )
                                    
                                    if deploy_result.get("success"):
                                        vercel_url = deploy_result.get("data", {}).get("url")
                                        steps.append(f"✅ Deployed to Vercel")
                                        data["vercel_url"] = vercel_url
                                        data["deployment_url"] = vercel_url
                                    else:
                                        steps.append(f"⚠️ Vercel deployment failed: {deploy_result.get('error')}")
                                        
                                except Exception as e:
                                    steps.append(f"⚠️ Vercel error: {str(e)}")
                                    logger.error(f"Vercel error: {e}", exc_info=True)
                            else:
                                steps.append("⚠️ Skipped Vercel (VERCEL_TOKEN not set)")
                            
                            # Success!
                            result = {
                                "success": True,
                                "completed": True,
                                "message": "Website created successfully!",
                                "steps": steps,
                                "data": data
                            }
                            
                            logger.info(f"🎉 Complete workflow finished!")
                        
                    except ImportError as e:
                        logger.error(f"Import error: {e}", exc_info=True)
                        result = {
                            "success": False,
                            "error": f"Failed to import Figma converter: {str(e)}",
                            "fatal": True
                        }
                    except Exception as e:
                        logger.error(f"Figma workflow error: {e}", exc_info=True)
                        result = {
                            "success": False,
                            "error": f"Workflow failed: {str(e)}",
                            "fatal": True
                        }
            
            # Document Generator
            elif name == "create_document_and_presentation":
                if not self.doc_gen:
                    result = {"success": False, "error": "Document Generator not configured", "fatal": True}
                else:
                    result = await self.doc_gen.create_document_and_presentation(
                        arguments["request"],
                        arguments.get("project_name"),
                        arguments.get("output_folder")
                    )
            
            # Screenshot tools
            elif name == "enable_screenshots":
                self.screenshot_enabled = True
                self._init_screenshot_doc()
                result = {"success": True, "completed": True, "message": "Screenshots enabled"}
            
            elif name == "disable_screenshots":
                self.screenshot_enabled = False
                doc_path = self._save_screenshot_doc()
                result = {"success": True, "completed": True, "message": f"Document saved: {doc_path}", "data": {"doc_path": doc_path}}
            
            elif name == "take_screenshot_and_save":
                if not self.browser.page:
                    result = {"success": False, "error": "Browser not started"}
                else:
                    if not self.screenshot_doc:
                        self._init_screenshot_doc()
                    screenshot_path = await self._take_and_save_screenshot(arguments.get("title", "Screenshot"))
                    result = {"success": True, "completed": True, "message": f"Screenshot saved: {screenshot_path}"}
            
            # ⭐ NEW: Code Analysis tools
            elif name == "analyze_code":
                if not self.code_analysis:
                    result = {"success": False, "error": "Code analysis not configured. AI client required.", "fatal": True}
                else:
                    result = await self.code_analysis.analyze_code(arguments["file_path"])

            elif name == "fix_bugs":
                if not self.code_analysis:
                    result = {"success": False, "error": "Code analysis not configured. AI client required.", "fatal": True}
                else:
                    result = await self.code_analysis.fix_bugs(
                        arguments["file_path"],
                        arguments.get("context")
                    )

            elif name == "generate_tests":
                if not self.code_analysis:
                    result = {"success": False, "error": "Code analysis not configured. AI client required.", "fatal": True}
                else:
                    result = await self.code_analysis.generate_tests(arguments["file_path"])

            elif name == "refactor_code":
                if not self.code_analysis:
                    result = {"success": False, "error": "Code analysis not configured. AI client required.", "fatal": True}
                else:
                    result = await self.code_analysis.refactor_code(arguments["file_path"])

            elif name == "document_code":
                if not self.code_analysis:
                    result = {"success": False, "error": "Code analysis not configured. AI client required.", "fatal": True}
                else:
                    result = await self.code_analysis.document_code(arguments["file_path"])

            # ⭐ NEW: Web Scraping tools
            elif name == "summarize_article":
                if not self.web_scraper:
                    result = {"success": False, "error": "Web scraper not configured. AI client required.", "fatal": True}
                else:
                    result = await self.web_scraper.summarize_article(arguments.get("url"))

            elif name == "scrape_table_to_csv":
                if not self.web_scraper:
                    result = {"success": False, "error": "Web scraper not configured.", "fatal": True}
                else:
                    result = await self.web_scraper.scrape_table_to_csv(
                        arguments.get("selector"),
                        arguments.get("filename", "scraped_table.csv")
                    )

            elif name == "research_competitors":
                if not self.web_scraper:
                    result = {"success": False, "error": "Web scraper not configured. AI client required.", "fatal": True}
                else:
                    result = await self.web_scraper.research_competitors(
                        arguments["topic"],
                        arguments.get("num_results", 5)
                    )
            
            # ⭐ NEW: Windows System Automation (after research_competitors handler)
            elif name == "transfer_files":
                result = await self._send_to_windows_context("transfer_files", {
                    "source": arguments["source"],
                    "destination": arguments["destination"],
                    "files": arguments.get("files", [])
                })

            elif name == "move_files":
                result = await self._send_to_windows_context("move_files", {
                    "source": arguments["source"],
                    "destination": arguments["destination"],
                    "files": arguments.get("files", [])
                })

            elif name == "install_app":
                result = await self._send_to_windows_context("install_app", {
                    "appName": arguments["app_name"]
                })

            elif name == "launch_app":
                result = await self._send_to_windows_context("launch_app", {
                    "appName": arguments["app_name"]
                })

            elif name == "close_app":
                result = await self._send_to_windows_context("close_app", {
                    "appName": arguments["app_name"]
                })

            elif name == "get_system_info":
                result = await self._send_to_windows_context("get_system_info", {})

            elif name == "get_running_apps":
                result = await self._send_to_windows_context("get_running_apps", {})

            elif name == "open_folder":
                result = await self._send_to_windows_context("open_folder", {
                    "path": arguments["path"]
                })

            elif name == "create_folder":
                result = await self._send_to_windows_context("create_folder", {
                    "path": arguments["path"]
                })
            
            else:
                result = {"success": False, "error": f"Unknown tool: {name}", "fatal": True}
            
            # Log usage
            self.db.log_tool_usage(name, result.get("success", False), 0)
            
            # CHECK FOR COMPLETION OR FATAL ERROR
            if result.get("completed") is True:
                logger.info(f"✅ Tool '{name}' marked as completed")
            elif result.get("fatal") is True:
                logger.error(f"❌ Tool '{name}' returned fatal error")
            
            return result
            
        except Exception as e:
            logger.error(f"Tool execution error: {e}", exc_info=True)
            return {"success": False, "error": str(e), "fatal": True}
    
    async def _send_to_windows_context(self, command: str, params: dict):
        """Send command to Windows context tracker via WebSocket"""
        try:
            uri = 'ws://localhost:8767'
            
            async with websockets.connect(uri) as ws:
                # Send command
                await ws.send(json.dumps({
                    "command": command,
                    "params": params
                }))
                
                # Wait for response
                response = await asyncio.wait_for(ws.recv(), timeout=30)
                result = json.loads(response)
                
                # Extract actual result
                if result.get('type') == 'command_result':
                    return result.get('result', {"success": False, "error": "No result"})
                
                return {"success": False, "error": "Invalid response format"}
                
        except websockets.exceptions.WebSocketException as e:
            logger.error(f"Windows context WebSocket error: {e}")
            return {
                "success": False, 
                "error": "Windows context tracker not running. Start: node extensions/windows-context/index.js",
                "fatal": True
            }
        except asyncio.TimeoutError:
            return {
                "success": False,
                "error": "Windows context tracker timeout",
                "fatal": True
            }
        except Exception as e:
            logger.error(f"Windows context error: {e}")
            return {
                "success": False,
                "error": f"Windows automation error: {str(e)}",
                "fatal": True
            }